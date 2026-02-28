# RemBot_IA_v3.py
# ✅ Modelo único (invariável) para prever:
#   (1) Semana2 a partir de Pre + Semana1
#   (2) Pós-teste a partir de Pre + Semana1 + Semana2
#   (3) Retenção a partir de Pre + Semana1 + Semana2 + Pós
#
# Estratégia: um único regressor com entrada fixa de 8 dims = [valores(4) + máscara(4)].
# • Valores ausentes são preenchidos com 0, e a máscara indica quais entradas existem.
# • No treino, cada aluno pode gerar até 3 amostras (targets distintos), dependendo das colunas disponíveis.

import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import random
import re
import math

# =========================
# Fonte global do sistema
# =========================
DEFAULT_FONT = ("Helvetica", 14)
DEFAULT_MONO = ("Courier", 12)

# 🔧 LIMIAR DO RELATÓRIO DE BAIXO RENDIMENTO (MUDE AQUI MANUALMENTE)
LIMIAR_BAIXO_RENDIMENTO = 40.0

# 🔧 LIMIAR PARA CONSIDERAR UMA QUESTÃO COMO "ERRO" (nota < limiar)
LIMIAR_ERRO_QUESTAO = 60.0

# =========================
# Lazy imports (carregam só quando precisa)
# =========================
np = None
pd = None
torch = None
nn = None
optim = None
plt = None
FigureCanvasTkAgg = None
RedePrevisao = None


def ensure_data_libs():
    """Carrega numpy/pandas sob demanda."""
    global np, pd
    if np is None:
        import numpy as _np
        np = _np
    if pd is None:
        import pandas as _pd
        pd = _pd


def ensure_plot_libs():
    """Carrega matplotlib + backend TkAgg sob demanda."""
    global plt, FigureCanvasTkAgg
    if plt is None or FigureCanvasTkAgg is None:
        import matplotlib.pyplot as _plt
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg as _FigureCanvasTkAgg
        plt = _plt
        FigureCanvasTkAgg = _FigureCanvasTkAgg


def ensure_torch_libs():
    """Carrega torch sob demanda."""
    global torch, nn, optim
    if torch is None:
        import torch as _torch
        import torch.nn as _nn
        import torch.optim as _optim
        torch = _torch
        nn = _nn
        optim = _optim


def get_model_class():
    """Define a classe do modelo só depois de torch estar disponível."""
    global RedePrevisao
    if RedePrevisao is None:
        ensure_torch_libs()

        class _RedePrevisao(nn.Module):
            # 8 entradas = 4 valores + 4 máscaras
            def __init__(self):
                super(_RedePrevisao, self).__init__()
                self.fc1 = nn.Linear(8, 32)
                self.fc2 = nn.Linear(32, 16)
                self.fc3 = nn.Linear(16, 1)
                self.relu = nn.ReLU()

            def forward(self, x):
                x = self.relu(self.fc1(x))
                x = self.relu(self.fc2(x))
                x = self.fc3(x)
                return x

        RedePrevisao = _RedePrevisao

    return RedePrevisao


# --- Tooltip leve para Tkinter ---
class Tooltip:
    def __init__(self, widget, text, delay=600):
        self.widget = widget
        self.text = text
        self.delay = delay
        self._id = None
        self._tip = None
        widget.bind("<Enter>", self._enter)
        widget.bind("<Leave>", self._leave)

    def _enter(self, _):
        self._schedule()

    def _leave(self, _):
        self._unschedule()
        self._hide()

    def _schedule(self):
        self._unschedule()
        self._id = self.widget.after(self.delay, self._show)

    def _unschedule(self):
        if self._id:
            self.widget.after_cancel(self._id)
            self._id = None

    def _show(self):
        if self._tip or not self.text:
            return
        x = (self.widget.winfo_rootx() + 20)
        y = (self.widget.winfo_rooty() + 20)
        self._tip = tw = tk.Toplevel(self.widget)
        tw.wm_overrideredirect(True)
        tw.wm_geometry(f"+{x}+{y}")
        frame = tk.Frame(tw, background="#ffffe0", borderwidth=1, relief="solid")
        label = tk.Label(
            frame,
            text=self.text,
            background="#ffffe0",
            justify="left",
            font=("Helvetica", 12),
            wraplength=420,
        )
        label.pack(padx=6, pady=4)
        frame.pack()

    def _hide(self):
        if self._tip:
            self._tip.destroy()
            self._tip = None


# ========================
# Utilitários
# ========================
def set_seed(seed=42):
    ensure_data_libs()
    ensure_torch_libs()
    random.seed(seed)
    np.random.seed(seed)
    torch.manual_seed(seed)
    torch.cuda.manual_seed_all(seed)
    torch.backends.cudnn.deterministic = True
    torch.backends.cudnn.benchmark = False


def r2_score_np(y_true, y_pred):
    ensure_data_libs()
    y_true = np.asarray(y_true, dtype=float).ravel()
    y_pred = np.asarray(y_pred, dtype=float).ravel()
    ss_res = np.sum((y_true - y_pred) ** 2)
    ss_tot = np.sum((y_true - np.mean(y_true)) ** 2)
    return 1.0 - ss_res / ss_tot if ss_tot > 0 else 0.0


def train_val_split_stratified(df, label_col="Grupo", test_size=0.2, seed=42):
    """Split estratificado por grupo (sem sklearn)."""
    ensure_data_libs()
    rng = np.random.default_rng(seed)
    idx_train, idx_val = [], []
    for _, df_g in df.groupby(label_col):
        perm = df_g.index.to_numpy()
        rng.shuffle(perm)
        n = len(perm)
        n_val = max(1, int(round(test_size * n)))
        idx_val.extend(perm[:n_val].tolist())
        idx_train.extend(perm[n_val:].tolist())
    return sorted(idx_train), sorted(idx_val)


# ========================
# Aplicação
# ========================
class EduPredict:
    def __init__(self, root):
        self.root = root
        self.root.title("🎓 RemBot IA - Análise e Previsão (Modelo Único Invariável)")
        self.root.geometry("1150x780")
        self.root.resizable(True, True)

        # Fonte global
        self.root.option_add("*Font", DEFAULT_FONT)

        # Estado
        self.dados = None
        self.modelo = None
        self.y_real = None
        self.y_previsto = None
        self.caminho_modelo = "modelo_salvo.pth"
        self.tema_escuro = False

        # Normalização (apenas do TREINO) — aplicada SOMENTE nos 4 valores (não na máscara)
        self.x_mean = None  # shape (4,)
        self.x_std = None   # shape (4,)
        self.usar_normalizacao = True

        # Normalização do ALVO por tarefa (Semana2/Pos_teste/Retencao) — calculada no treino
        # Formato: {task: {'mean': float, 'std': float}}
        self.y_stats = {}

        # Calibração (mantida; aplicada ao output)
        self.calib = {"global": None, "group_alpha": {}, "group_lin": {}}
        self.calib_method = "lin_group"  # "none" | "alpha_global" | "alpha_group" | "lin_group"

        # Split por ALUNO (para não vazar amostras do mesmo aluno entre treino/val)
        self.idx_train = None
        self.idx_val = None
        self.test_size = 0.2
        self.seed = 42

        # Early stopping
        self.patience = 50
        self.min_delta = 1e-5

        # Visualização Treino/Val
        self.view_subset = "val"
        self.view_subset_var = tk.StringVar(value="val")

        # Cache de métricas/predições (agora inclui por tarefa)
        self.cache = {
            "train": {"y_true": None, "y_pred": None, "mse": None, "r2": None, "task": None, "by_task": {}},
            "val": {"y_true": None, "y_pred": None, "mse": None, "r2": None, "task": None, "by_task": {}},
        }

        # Colunas disponíveis para curva (dataset parcial)
        self.colunas_disponiveis_curva = []  # lista de (col, label)

        # Estilo (leve)
        self.style = ttk.Style()
        self.style.theme_use("clam")
        self.style.configure(".", font=DEFAULT_FONT)
        self.style.configure("TButton", font=DEFAULT_FONT, padding=6)
        self.style.configure("TLabel", font=DEFAULT_FONT)
        self.style.configure("Header.TLabel", font=("Helvetica", 18, "bold"))
        self.style.configure("TNotebook.Tab", font=("Helvetica", 13, "bold"))

        self.criar_interface()
        self.criar_menu()

        # ✅ Splash/feedback inicial (sem importar libs pesadas)
        self.status_var.set("Inicializando interface...")
        self.root.after(50, self._pos_inicializacao)

    def _pos_inicializacao(self):
        self.status_var.set("✅ Pronto. Carregue os dados para começar...")

    # =======================
    # Identificação do aluno
    # =======================
    def _get_ident_aluno(self, idx: int) -> str:
        if self.dados is None:
            return f"Aluno {idx+1}"

        row = self.dados.iloc[idx]

        for col in ["Matricula", "Matrícula"]:
            if col in self.dados.columns:
                v = row.get(col, None)
                if v is not None and str(v).strip() != "" and str(v).lower() != "nan":
                    return str(v).strip()

        if "Aluno" in self.dados.columns:
            v = row.get("Aluno", None)
            if v is not None and str(v).strip() != "" and str(v).lower() != "nan":
                return str(v).strip()

        return f"Aluno {idx+1}"

    # =======================
    # Detecta colunas disponíveis para curva
    # =======================
    def _detectar_colunas_curva(self):
        if self.dados is None:
            self.colunas_disponiveis_curva = []
            return

        ordem = [
            ("Pre_teste", "Pré"),
            ("Semana1", "1 Sem"),
            ("Semana2", "2 Sem"),
            ("Pos_teste", "Pós"),
        ]
        disp = []
        for c, lab in ordem:
            if c in self.dados.columns and self.dados[c].notna().any():
                disp.append((c, lab))

        if "Retencao" in self.dados.columns and self.dados["Retencao"].notna().any():
            disp.append(("Retencao", "Retenção"))

        self.colunas_disponiveis_curva = disp

    def _tem_base_pre_sem1(self):
        if self.dados is None:
            return False
        return (
            ("Pre_teste" in self.dados.columns and self.dados["Pre_teste"].notna().any())
            and ("Semana1" in self.dados.columns and self.dados["Semana1"].notna().any())
        )

    def _tem_alvo_sem2(self):
        return self.dados is not None and ("Semana2" in self.dados.columns) and self.dados["Semana2"].notna().any()

    def _tem_alvo_pos(self):
        return self.dados is not None and ("Pos_teste" in self.dados.columns) and self.dados["Pos_teste"].notna().any()

    def _tem_alvo_retencao(self):
        return self.dados is not None and ("Retencao" in self.dados.columns) and self.dados["Retencao"].notna().any()

    def _tem_dados_para_treino_invariavel(self):
        """Condição mínima para treinar o modelo único:
        - Precisa ter Pre_teste e Semana1 (para montar entrada mínima)
        - E precisa ter pelo menos 1 alvo disponível (Semana2 ou Pos_teste ou Retencao)
        """
        if self.dados is None:
            return False
        if not self._tem_base_pre_sem1():
            return False
        return self._tem_alvo_sem2() or self._tem_alvo_pos() or self._tem_alvo_retencao()

    # =======================
    # Classificação (só no resumo por grupo)
    # =======================
    @staticmethod
    def _classe_por_nota(v: float) -> str:
        if v < 40:
            return "Risco crítico (<40)"
        if v < 50:
            return "Recuperação (<50)"
        if v < 60:
            return "Risco acadêmico (<60)"
        if v < 70:
            return "Excelência mínima (<70)"
        return "OK (>=70)"

    # =======================
    # Menu
    # =======================
    def criar_menu(self):
        menubar = tk.Menu(self.root)

        menu_ajuda = tk.Menu(menubar, tearoff=0)
        menu_ajuda.add_command(label="Manual do Usuário", command=self.manual)
        menu_ajuda.add_command(label="Sobre", command=self.sobre)
        menubar.add_cascade(label="❓ Ajuda", menu=menu_ajuda)

        menu_cfg = tk.Menu(menubar, tearoff=0)

        sub_norm = tk.Menu(menu_cfg, tearoff=0)
        sub_norm.add_radiobutton(label="Usar normalização (recomendado)", command=lambda: self.set_norm(True))
        sub_norm.add_radiobutton(label="Sem normalização", command=lambda: self.set_norm(False))
        menu_cfg.add_cascade(label="Normalização", menu=sub_norm)

        sub_cal = tk.Menu(menu_cfg, tearoff=0)
        sub_cal.add_radiobutton(label="Sem calibração", command=lambda: self.set_cal("none"))
        sub_cal.add_radiobutton(label="α global (médias)", command=lambda: self.set_cal("alpha_global"))
        sub_cal.add_radiobutton(label="α por grupo", command=lambda: self.set_cal("alpha_group"))
        sub_cal.add_radiobutton(label="Regressão por grupo (a, b)", command=lambda: self.set_cal("lin_group"))
        menu_cfg.add_cascade(label="Calibração", menu=sub_cal)

        sub_view = tk.Menu(menu_cfg, tearoff=0)
        sub_view.add_radiobutton(
            label="Mostrar Validação",
            variable=self.view_subset_var,
            value="val",
            command=lambda: self.set_view_subset("val"),
        )
        sub_view.add_radiobutton(
            label="Mostrar Treino",
            variable=self.view_subset_var,
            value="train",
            command=lambda: self.set_view_subset("train"),
        )
        menu_cfg.add_cascade(label="Visualização", menu=sub_view)

        menubar.add_cascade(label="⚙️ Configurações", menu=menu_cfg)
        self.root.config(menu=menubar)

    def manual(self):
        messagebox.showinfo(
            "Manual",
            "1) Carregue o CSV\n"
            "2) A aba Gráficos mostra a curva mesmo com dataset parcial\n"
            "3) Treino IA (modelo único) exige: Pre_teste + Semana1 e pelo menos 1 alvo preenchido:\n"
            "   • Semana2 (para prever Semana2)\n"
            "   • Pos_teste (para prever Pós)\n"
            "   • Retencao (para prever Retenção)\n"
            "4) Modelo único é invariável: aceita 2, 3 ou 4 entradas (via máscara)\n"
            "5) Identificação usa Matricula/Matrícula ou Aluno do CSV."
        )

    def sobre(self):
        messagebox.showinfo(
            "EduPredict",
            "Versão invariável (um único modelo)\n"
            "• Entrada variável (2/3/4 notas) com máscara\n"
            "• Treino multitarefa: prevê Semana2, Pós e Retenção\n"
            "• Inicialização rápida (lazy imports)\n"
            "• Curva e tabelas por tema/nível (Q1..Qn)\n"
        )

    def set_view_subset(self, subset: str):
        self.view_subset = subset
        if self.cache[subset]["y_true"] is None:
            messagebox.showwarning("Visualização", "Ainda não há métricas para esta partição. Treine o modelo primeiro.")
            return
        self.update_view_data_and_plot()

    def update_view_data_and_plot(self):
        part = self.view_subset
        y_true = self.cache[part]["y_true"]
        y_pred = self.cache[part]["y_pred"]
        mse = self.cache[part]["mse"]
        r2 = self.cache[part]["r2"]
        self.y_real = y_true
        self.y_previsto = y_pred
        titulo = "(Validação)" if part == "val" else "(Treino)"
        self.mostrar_resultados_modelo(mse, r2, titulo_subset=titulo)

    def set_norm(self, flag):
        self.usar_normalizacao = flag
        messagebox.showinfo("Normalização", f"Normalização {'ativada' if flag else 'desativada'}. Re-treine o modelo.")

    def set_cal(self, metodo):
        self.calib_method = metodo
        messagebox.showinfo("Calibração", f"Método de calibração definido como: {metodo}")

    # =======================
    # Interface
    # =======================
    def criar_interface(self):
        notebook = ttk.Notebook(self.root)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        self.tab_inicio = ttk.Frame(notebook)
        self.tab_graficos = ttk.Frame(notebook)
        self.tab_modelo = ttk.Frame(notebook)
        self.tab_relatorio = ttk.Frame(notebook)
        self.tab_previsao = ttk.Frame(notebook)

        notebook.add(self.tab_inicio, text="🏠 Início")
        notebook.add(self.tab_graficos, text="📊 Gráficos")
        notebook.add(self.tab_modelo, text="🤖 Modelo")
        notebook.add(self.tab_relatorio, text="📄 Relatório")
        notebook.add(self.tab_previsao, text="🔮 Previsão Individual")

        self.criar_aba_inicio()
        self.criar_aba_graficos()
        self.criar_aba_modelo()
        self.criar_aba_relatorio()
        self.criar_aba_previsao()

        self.status_var = tk.StringVar(value="Carregue os dados para começar...")
        self.status_bar = ttk.Label(self.root, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W)
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)

    def criar_aba_inicio(self):
        frame = self.tab_inicio
        ttk.Label(frame, text="Análise e Previsão (Modelo Único Invariável)", style="Header.TLabel").pack(pady=20)

        desc = (
            "✅ Modelo único e invariável (aceita 2/3/4 entradas)\n"
            "✅ Previsões: Semana2, Pós e Retenção (dependendo do que você informar)\n"
            "✅ Curva de aprendizagem aparece mesmo com dataset parcial\n"
            "✅ Tabelas: erros por Tema e por Nível cognitivo (se Q1..Qn existir)\n"
        )
        tk.Label(frame, text=desc, justify=tk.CENTER, font=DEFAULT_FONT, fg="gray").pack(pady=20)

        btn_frame = ttk.Frame(frame)
        btn_frame.pack(pady=30)

        self.btn_carregar = ttk.Button(btn_frame, text="📁 Carregar Dados (CSV)", width=25, command=self.carregar_csv)
        self.btn_carregar.pack(pady=8)
        Tooltip(self.btn_carregar, "CSV deve conter pelo menos a coluna 'Grupo'. Matricula/Aluno é recomendado.")

        tema_btn = ttk.Button(btn_frame, text="🌓 Alternar Tema", command=self.alternar_tema)
        tema_btn.pack(pady=8)

    def alternar_tema(self):
        if self.tema_escuro:
            self.style.theme_use("clam")
            self.root.configure(bg="white")
        else:
            self.style.theme_use("alt")
            self.root.configure(bg="#2e2e2e")
        self.tema_escuro = not self.tema_escuro


    def _sort_tree(self, tree: ttk.Treeview, col: str, reverse: bool):
        """Ordena Treeview por coluna (texto/número)."""
        try:
            data = [(tree.set(k, col), k) for k in tree.get_children("")]
            def to_key(v):
                v = str(v).strip()
                if v.upper() in ("SN", "", "NAN", "NONE"):
                    return (1e18,)
                # tenta número
                try:
                    return (float(v.replace(",", ".")),)
                except Exception:
                    return (v.lower(),)
            data.sort(key=lambda x: to_key(x[0]), reverse=reverse)
            for idx, (_, k) in enumerate(data):
                tree.move(k, "", idx)
            tree.heading(col, command=lambda: self._sort_tree(tree, col, not reverse))
        except Exception:
            pass

    def criar_aba_graficos(self):
        frame = self.tab_graficos
        ttk.Label(frame, text="Visualização", style="Header.TLabel").pack(pady=10)

        container = ttk.Frame(frame)
        container.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        container.columnconfigure(0, weight=2)
        container.columnconfigure(1, weight=1)
        container.columnconfigure(2, weight=1)
        container.rowconfigure(0, weight=1)

        # 1) Curva de aprendizagem (média por grupo)
        self.lf_curva = ttk.LabelFrame(container, text="Curva de Aprendizagem (média por grupo)", padding=10)
        self.lf_curva.grid(row=0, column=0, sticky="nsew", padx=(0, 10), pady=0)
        self.frame_grafico1 = ttk.Frame(self.lf_curva)
        self.frame_grafico1.pack(fill=tk.BOTH, expand=True)
        ttk.Label(self.frame_grafico1, text="Carregue um CSV para exibir a curva.").pack(pady=25)

        # 2) Top temas com mais incidência de erro
        self.lf_temas = ttk.LabelFrame(container, text="Top 7 temas com mais incidência de erro (dispersão)", padding=10)
        self.lf_temas.grid(row=0, column=1, sticky="nsew", padx=(0, 10), pady=0)

        cols = ("Tema", "Erro(%)", "Total")
        self.tree_temas = ttk.Treeview(self.lf_temas, columns=cols, show="headings", height=14)
        for c, w in zip(cols, (220, 70, 70)):
            self.tree_temas.heading(c, text=c)
            self.tree_temas.column(c, width=w, anchor="w" if c == "Tema" else "center", stretch=True)

        sb1y = ttk.Scrollbar(self.lf_temas, orient="vertical", command=self.tree_temas.yview)
        self.tree_temas.configure(yscrollcommand=sb1y.set)
        self.tree_temas.grid(row=0, column=0, sticky="nsew")
        sb1y.grid(row=0, column=1, sticky="ns")
        self.lf_temas.rowconfigure(0, weight=1)
        self.lf_temas.columnconfigure(0, weight=1)

        self.lbl_hint_tema = ttk.Label(
            self.lf_temas,
            text=f"Inclua colunas Q1..Qn (0/1, 0..10, 0..100) e Tema_Q1..Tema_Qn. Erro(%) = 100*(1 - nota_norm).",
            wraplength=260,
        )
        self.lbl_hint_tema.grid(row=1, column=0, columnspan=2, sticky="w", pady=(8, 0))

        # 3) Erros por nível cognitivo
        self.lf_niveis = ttk.LabelFrame(container, text="Erros por nível cognitivo (dispersão)", padding=10)
        self.lf_niveis.grid(row=0, column=2, sticky="nsew", padx=0, pady=0)

        cols2 = ("Nível", "Erros", "Total")
        self.tree_niveis = ttk.Treeview(self.lf_niveis, columns=cols2, show="headings", height=14)
        for c, w in zip(cols2, (180, 70, 70)):
            self.tree_niveis.heading(c, text=c)
            self.tree_niveis.column(c, width=w, anchor="w" if c == "Nível" else "center", stretch=True)

        sb2y = ttk.Scrollbar(self.lf_niveis, orient="vertical", command=self.tree_niveis.yview)
        self.tree_niveis.configure(yscrollcommand=sb2y.set)
        self.tree_niveis.grid(row=0, column=0, sticky="nsew")
        sb2y.grid(row=0, column=1, sticky="ns")
        self.lf_niveis.rowconfigure(0, weight=1)
        self.lf_niveis.columnconfigure(0, weight=1)

        self.lbl_hint_nivel = ttk.Label(
            self.lf_niveis,
            text=f"Inclua colunas Q1..Qn (0/1, 0..10, 0..100) e Nivel_Q1..Nivel_Qn. Erro(%) = 100*(1 - nota_norm).",
            wraplength=260,
        )
        self.lbl_hint_nivel.grid(row=1, column=0, columnspan=2, sticky="w", pady=(8, 0))

    def criar_aba_modelo(self):
        frame = self.tab_modelo
        ttk.Label(frame, text="Treinamento do Modelo Único", style="Header.TLabel").pack(pady=10)

        self.info_modelo = tk.Text(frame, height=12, width=92, font=DEFAULT_MONO, state=tk.DISABLED)
        self.info_modelo.pack(pady=10, padx=20)

        btn_frame = ttk.Frame(frame)
        btn_frame.pack(pady=5)

        self.btn_treinar = ttk.Button(btn_frame, text="🚀 Treinar Modelo", width=20, command=self.treinar, state="disabled")
        self.btn_treinar.pack(side=tk.LEFT, padx=5)

        self.btn_salvar = ttk.Button(btn_frame, text="💾 Salvar Modelo", width=20, command=self.salvar_modelo, state="disabled")
        self.btn_salvar.pack(side=tk.LEFT, padx=5)

        self.btn_carregar_modelo = ttk.Button(btn_frame, text="📂 Carregar Modelo", width=20, command=self.carregar_modelo)
        self.btn_carregar_modelo.pack(side=tk.LEFT, padx=5)

        self.btn_atualizar = ttk.Button(
            btn_frame, text="🔄 Atualizar com Novos Dados", width=25, command=self.atualizar_modelo, state="disabled"
        )
        self.btn_atualizar.pack(side=tk.LEFT, padx=5)

        Tooltip(self.btn_treinar, "Treina modelo único com máscara. Requer Pre_teste+Semana1 e ao menos 1 alvo (Semana2/Pós/Retenção).")
        Tooltip(self.btn_atualizar, "Refaz o treino completo com os dados atuais.")
        Tooltip(self.btn_salvar, "Salva apenas os pesos do modelo.")
        Tooltip(self.btn_carregar_modelo, "Carrega pesos salvos (.pth).")

        self.frame_grafico2 = ttk.Frame(frame)
        self.frame_grafico2.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)

    def criar_aba_relatorio(self):
            frame = self.tab_relatorio
            ttk.Label(frame, text="Relatório Automático", style="Header.TLabel").pack(pady=10)

            self.lbl_relatorio_prev = ttk.Label(frame, text="PREVISÕES — (carregue dados e treine/carregue um modelo)", font=DEFAULT_FONT)
            self.lbl_relatorio_prev.pack(anchor="w", padx=20, pady=(0, 6))

            text_frame = ttk.Frame(frame)
            text_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=(0, 10))
            text_frame.rowconfigure(0, weight=1)
            text_frame.columnconfigure(0, weight=1)

            self.txt_relatorio = tk.Text(text_frame, wrap="none", font=("Consolas", 11))
            self.txt_relatorio.grid(row=0, column=0, sticky="nsew")

            sb_y = ttk.Scrollbar(text_frame, orient="vertical", command=self.txt_relatorio.yview)
            sb_x = ttk.Scrollbar(text_frame, orient="horizontal", command=self.txt_relatorio.xview)
            sb_y.grid(row=0, column=1, sticky="ns")
            sb_x.grid(row=1, column=0, sticky="ew")

            self.txt_relatorio.configure(yscrollcommand=sb_y.set, xscrollcommand=sb_x.set)

            btn_frame = ttk.Frame(frame)
            btn_frame.pack(pady=10)

            ttk.Button(btn_frame, text="📝 Gerar Relatório Word", command=self.exportar_word).grid(row=0, column=0, padx=10)
            ttk.Button(btn_frame, text="📊 Exportar Excel", command=self.exportar_excel).grid(row=0, column=1, padx=10)
            ttk.Button(btn_frame, text="💾 Salvar Gráfico", command=self.salvar_grafico).grid(row=0, column=2, padx=10)

    def criar_aba_previsao(self):
        frame = self.tab_previsao
        ttk.Label(frame, text="Previsão Individual (Modelo Único)", style="Header.TLabel").pack(pady=15)

        instrucoes = (
            "Digite as notas disponíveis. O modelo escolhe automaticamente o alvo:\n"
            "• Pre + Semana1  -> prevê Semana2\n"
            "• Pre + Semana1 + Semana2 -> prevê Pós\n"
            "• Pre + Semana1 + Semana2 + Pós -> prevê Retenção\n"
            "⚠ Use valores 0–100. Campos vazios são aceitos (desde que a sequência faça sentido)."
        )
        tk.Label(frame, text=instrucoes, justify=tk.CENTER, font=DEFAULT_FONT, fg="gray").pack(pady=10)

        entrada_frame = ttk.LabelFrame(frame, text="Dados do Aluno", padding=15)
        entrada_frame.pack(padx=40, pady=20, fill="x")

        self.var_grupo = tk.StringVar(value="Tradicional")
        self.var_pre = tk.StringVar()
        self.var_sem1 = tk.StringVar()
        self.var_sem2 = tk.StringVar()
        self.var_pos = tk.StringVar()

        ttk.Label(entrada_frame, text="Grupo:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        ttk.Combobox(
            entrada_frame,
            textvariable=self.var_grupo,
            values=["Tradicional", "Digital", "RemBot"],
            state="readonly",
            width=15,
        ).grid(row=0, column=1, padx=5, pady=5)

        ttk.Label(entrada_frame, text="Pré-teste:").grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
        ttk.Entry(entrada_frame, textvariable=self.var_pre, width=15).grid(row=1, column=1, padx=5, pady=5)

        ttk.Label(entrada_frame, text="Semana 1:").grid(row=2, column=0, sticky=tk.W, padx=5, pady=5)
        ttk.Entry(entrada_frame, textvariable=self.var_sem1, width=15).grid(row=2, column=1, padx=5, pady=5)

        ttk.Label(entrada_frame, text="Semana 2:").grid(row=3, column=0, sticky=tk.W, padx=5, pady=5)
        ttk.Entry(entrada_frame, textvariable=self.var_sem2, width=15).grid(row=3, column=1, padx=5, pady=5)

        ttk.Label(entrada_frame, text="Pós-teste:").grid(row=4, column=0, sticky=tk.W, padx=5, pady=5)
        ttk.Entry(entrada_frame, textvariable=self.var_pos, width=15).grid(row=4, column=1, padx=5, pady=5)

        self.btn_prever = ttk.Button(frame, text="🔮 Prever (alvo automático)", command=self.prever_individual, state="disabled")
        self.btn_prever.pack(pady=15)

        self.resultado_var = tk.StringVar(value="Aguardando previsão...")
        ttk.Label(frame, textvariable=self.resultado_var, font=("Helvetica", 16, "bold"), foreground="blue").pack(pady=10)

        self.dica_label = ttk.Label(frame, text="💡 Carregue um modelo para habilitar a previsão", foreground="red")
        self.dica_label.pack(pady=5)

    # ========================
    # Plot curva (dataset parcial)
    # ========================
    def plotar_curvas_medias(self):
        if self.dados is None:
            return

        ensure_plot_libs()
        ensure_data_libs()

        self._detectar_colunas_curva()
        if len(self.colunas_disponiveis_curva) == 0:
            for widget in self.frame_grafico1.winfo_children():
                widget.destroy()
            ttk.Label(self.frame_grafico1, text="Nenhuma coluna válida encontrada para curva.").pack(pady=30)
            return

        plt.close("all")
        momentos = [c for c, _ in self.colunas_disponiveis_curva]
        labels = [l for _, l in self.colunas_disponiveis_curva]

        fig, ax = plt.subplots(figsize=(9, 5))
        cores_fixas = {"Tradicional": "#1f77b4", "Digital": "#2ca02c", "RemBot": "#ff7f0e"}
        cores_aleatorias = ["#d62728", "#9467bd", "#8c564b", "#e377c2"]

        grupos = list(self.dados["Grupo"].dropna().unique())
        for i, grupo in enumerate(grupos):
            subset = self.dados[self.dados["Grupo"] == grupo]
            medias = []
            for m in momentos:
                medias.append(float(subset[m].mean()) if subset[m].notna().any() else np.nan)

            cor = cores_fixas.get(grupo, cores_aleatorias[i % len(cores_aleatorias)])
            ax.plot(labels, medias, marker="o", label=grupo, color=cor, linewidth=2.5)

        ax.set_title("Curva de Aprendizagem Média por Grupo", fontsize=16, fontweight="bold")
        ax.set_ylabel("Nota Média", fontsize=14)
        ax.set_xlabel("Momento da Avaliação", fontsize=14)
        ax.legend(title="Método de Ensino", fontsize=12)
        ax.grid(True, alpha=0.3)

        vals = []
        for m in momentos:
            if m in self.dados.columns:
                vals += self.dados[m].dropna().tolist()
        if vals:
            lo = max(0, float(np.min(vals)) - 5)
            hi = min(100, float(np.max(vals)) + 5)
            ax.set_ylim(lo, hi)
        else:
            ax.set_ylim(0, 100)

        for widget in self.frame_grafico1.winfo_children():
            widget.destroy()

        canvas = FigureCanvasTkAgg(fig, self.frame_grafico1)
        canvas.draw()
        canvas.get_tk_widget().pack(pady=10)
        self.canvas1 = canvas

    # ========================
    # Erros por Tema / Nível (Q1..Qn)
    # ========================
    def _detectar_colunas_questoes(self):
        if self.dados is None:
            return []
        qs = [c for c in self.dados.columns if re.fullmatch(r"Q\d+", str(c).strip())]
        return sorted(qs, key=lambda x: int(str(x)[1:]))

    def _limpar_tree(self, tree):
        for iid in tree.get_children():
            tree.delete(iid)


    def _inferir_escala_questao(self, serie: "pd.Series") -> float:
        """
        Infere a escala de uma questão para normalização em [0,1].

        Suporta:
          - binário 0/1
          - percentuais 0..100
          - notas contínuas (ex.: 0..10, 0..5)
        """
        ensure_data_libs()
        s = pd.to_numeric(serie, errors="coerce").dropna()
        if s.empty:
            return 1.0

        mx = float(s.max())

        # Amostra de valores únicos (evita custo alto)
        try:
            sample = s.values[: min(len(s), 500)]
            uniq = set(np.unique(sample))
        except Exception:
            uniq = set()

        # Heurísticas de escala
        if mx <= 1.0:
            return 1.0

        # Casos discretos pequenos
        if mx <= 2.0 and uniq and uniq.issubset({0.0, 1.0, 2.0}):
            return 1.0 if uniq.issubset({0.0, 1.0}) else mx

        if mx <= 10.0:
            # Se estiver claramente em 0..10, usa 10
            return 10.0

        if mx <= 100.0:
            return 100.0

        # Fallback: assume escala = máximo observado
        return mx

    def _calc_erros_agrupado(self, modo="tema"):
        """
        Calcula erros agregados por Tema ou por Nível cognitivo a partir de Q1..Qn.

        Dataset esperado:
            - Q1..Qn: desempenho por questão (pode ser 0/1, 0..10, 0..100, contínuo)
            - Tema_Q1..Tema_Qn (ou Nivel_Q1..Nivel_Qn)
        Definição de erro (não binária):
            score_norm = clip(nota / escala, 0..1)
            erro_norm  = 1 - score_norm
        O relatório exibe:
            - Erro(%) = 100 * média(erro_norm) no grupo (Tema/Nível)
            - Total = nº de respostas válidas consideradas
        """
        ensure_data_libs()

        qs = self._detectar_colunas_questoes()
        if not qs:
            return []

        map_prefix = "Tema_" if modo == "tema" else "Nivel_"
        if not any((f"{map_prefix}{q}" in self.dados.columns) for q in qs):
            # se não tiver mapeamento, ainda assim gera fallback por índice (Tema i / CPB/CPM/CAP)
            pass

        erro_sum = {}
        tot = {}

        for q in qs:
            notas_raw = pd.to_numeric(self.dados[q], errors="coerce")
            escala = float(self._inferir_escala_questao(notas_raw))
            score_norm = (notas_raw / escala).clip(lower=0.0, upper=1.0)
            err = (1.0 - score_norm)

            key_col = f"{map_prefix}{q}"
            if key_col in self.dados.columns:
                keys = self.dados[key_col].astype(str).fillna("").replace("nan", "").str.strip()
            else:
                keys = pd.Series([""] * len(self.dados))

            # fallback automático se o mapeamento estiver inconsistente (muitos valores únicos por aluno)
            try:
                uniq_ratio = float(keys.nunique(dropna=True)) / max(1, len(keys))
            except Exception:
                uniq_ratio = 1.0

            qnum = int(str(q)[1:])
            if (key_col not in self.dados.columns) or (uniq_ratio > 0.80):
                if modo == "tema":
                    keys = pd.Series([f"Tema {qnum}"] * len(keys))
                else:
                    # mapeamento simples (ajuste livre se quiser)
                    if qnum <= 4:
                        nivel = "CPB"
                    elif qnum <= 7:
                        nivel = "CPM"
                    else:
                        nivel = "CAP"
                    keys = pd.Series([nivel] * len(keys))

            for k, e, nota in zip(keys, err, notas_raw):
                if (not k) or pd.isna(nota):
                    continue
                tot[k] = tot.get(k, 0) + 1
                erro_sum[k] = erro_sum.get(k, 0.0) + float(e)

        rows = []
        for k in tot.keys():
            e = float(erro_sum.get(k, 0.0))
            t = int(tot[k])
            rows.append((k, e, t))

        rows.sort(key=lambda r: ((r[1] / r[2]) if r[2] else 0.0, r[1]), reverse=True)
        return rows

    def atualizar_tabelas_erros(self):
        if not hasattr(self, "tree_temas") or not hasattr(self, "tree_niveis"):
            return

        ensure_data_libs()

        self._limpar_tree(self.tree_temas)
        temas = self._calc_erros_agrupado("tema")
        if temas:
            self.lbl_hint_tema.config(text="")
            for k, e, t in temas[:7]:
                p = (100.0 * float(e) / float(t)) if t else 0.0
                self.tree_temas.insert("", "end", values=(k, f"{p:.1f}%", t))
        else:
            self.lbl_hint_tema.config(text=f"Inclua colunas Q1..Qn (0/1, 0..10, 0..100) e Tema_Q1..Tema_Qn. Erro(%) = 100*(1 - nota_norm).")

        self._limpar_tree(self.tree_niveis)
        niveis = self._calc_erros_agrupado("nivel")
        if niveis:
            self.lbl_hint_nivel.config(text="")
            for k, e, t in niveis:
                p = (100.0 * float(e) / float(t)) if t else 0.0
                self.tree_niveis.insert("", "end", values=(k, f"{p:.1f}%", t))
        else:
            self.lbl_hint_nivel.config(text=f"Inclua colunas Q1..Qn (0/1, 0..10, 0..100) e Nivel_Q1..Nivel_Qn. Erro(%) = 100*(1 - nota_norm).")

    # ========================
    # CSV
    # ========================
    def carregar_csv(self):
        caminho = filedialog.askopenfilename(title="Selecione o CSV", filetypes=[("CSV files", "*.csv")])
        if not caminho:
            return

        self.status_var.set("Carregando dataset...")

        try:
            ensure_data_libs()
            self.dados = pd.read_csv(caminho)

            if "Grupo" not in self.dados.columns:
                raise ValueError("Coluna obrigatória ausente: Grupo")

            # Auto-suporte a dataset por questões (Q1..Qn)
            # IMPORTANTE: Q1..Qn são usados apenas para tabelas/gráficos de erros (Tema/Nível).
            # NÃO derivamos automaticamente Semana2/Pós/Retenção a partir das questões, para não
            # "completar" datasets de teste (ex.: Pre+Sem1 com questões) e gerar preenchimentos indevidos.
            qs = [c for c in self.dados.columns if re.fullmatch(r"Q\d+", str(c).strip())]
            if qs:
                for q in qs:
                    self.dados[q] = pd.to_numeric(self.dados[q], errors="coerce")
                if "Retencao" not in self.dados.columns:
                    self.dados["Retencao"] = np.nan

            # Converte momentos para numérico se existirem
            for c in ["Pre_teste", "Semana1", "Semana2", "Pos_teste", "Retencao"]:
                if c in self.dados.columns:
                    self.dados[c] = pd.to_numeric(self.dados[c], errors="coerce")

            total = len(self.dados)
            grupos = self.dados["Grupo"].dropna().unique()
            self.status_var.set(f"✅ Dados carregados: {total} alunos, {len(grupos)} grupos")

            # plota curva
            self.plotar_curvas_medias()

            # atualiza tabelas de erro (Tema/Nível)
            self.atualizar_tabelas_erros()

            pode_treinar = self._tem_dados_para_treino_invariavel()
            self.btn_treinar.config(state="normal" if pode_treinar else "disabled")
            self.btn_atualizar.config(state="normal" if (self.modelo and pode_treinar) else "disabled")

            # Feedback ao usuário (quais alvos existem)
            msg = f"✅ Dataset carregado!\nAlunos: {total}\nGrupos: {len(grupos)}\n\nAlvos disponíveis:\n"
            msg += f"• Semana2: {'SIM' if self._tem_alvo_sem2() else 'NÃO'}\n"
            msg += f"• Pós-teste: {'SIM' if self._tem_alvo_pos() else 'NÃO'}\n"
            msg += f"• Retenção: {'SIM' if self._tem_alvo_retencao() else 'NÃO'}\n"
            if not self._tem_base_pre_sem1():
                msg += "\n⚠ Falta base mínima: Pre_teste e/ou Semana1."
            messagebox.showinfo("Sucesso", msg)

            self.gerar_relatorio()
            self.habilitar_previsao_individual()

        except Exception as e:
            self.status_var.set(f"❌ Erro: {str(e)}")
            messagebox.showerror("Erro", f"Erro ao carregar CSV:\n{str(e)}")

    # ========================
    # Dataset invariável (valores + máscara)
    # ========================
    @staticmethod
    def _mask_from_present(present_flags):
        # present_flags: list[bool] len=4
        return [1.0 if p else 0.0 for p in present_flags]

    def _build_samples_for_indices(self, idxs):
        """
        Constrói amostras (X,y) a partir de uma lista de índices de alunos.
        Para cada aluno, cria até 3 amostras:
          - X=[Pre,S1,0,0] + mask=[1,1,0,0] -> y=Semana2 (se existir)
          - X=[Pre,S1,S2,0] + mask=[1,1,1,0] -> y=Pos (se existir e S2 existir)
          - X=[Pre,S1,S2,Pos] + mask=[1,1,1,1] -> y=Retencao (se existir e S2 e Pos existirem)
        Retorna:
          X (n,8), y(n,1), task(list[str]), groups(list[str]), aluno_idx(list[int])
        """
        ensure_data_libs()

        X_list, y_list, task_list, g_list, aidx_list = [], [], [], [], []

        for i in idxs:
            row = self.dados.iloc[i]
            grupo = str(row.get("Grupo", ""))

            pre = row.get("Pre_teste", np.nan)
            s1 = row.get("Semana1", np.nan)
            s2 = row.get("Semana2", np.nan)
            pos = row.get("Pos_teste", np.nan)
            ret = row.get("Retencao", np.nan)

            # base mínima
            if pd.isna(pre) or pd.isna(s1):
                continue

            pre = float(pre)
            s1 = float(s1)

            # 1) prever Semana2
            if "Semana2" in self.dados.columns and pd.notna(s2):
                s2f = float(s2)
                vals = [pre, s1, 0.0, 0.0]
                msk = self._mask_from_present([True, True, False, False])
                X_list.append(vals + msk)
                y_list.append([s2f])
                task_list.append("Semana2")
                g_list.append(grupo)
                aidx_list.append(i)

            # 2) prever Pós (precisa de Semana2)
            if "Pos_teste" in self.dados.columns and pd.notna(pos) and ("Semana2" in self.dados.columns) and pd.notna(s2):
                s2f = float(s2)
                posf = float(pos)
                vals = [pre, s1, s2f, 0.0]
                msk = self._mask_from_present([True, True, True, False])
                X_list.append(vals + msk)
                y_list.append([posf])
                task_list.append("Pos_teste")
                g_list.append(grupo)
                aidx_list.append(i)

            # 3) prever Retenção (precisa de Semana2 e Pós)
            if "Retencao" in self.dados.columns and pd.notna(ret) and ("Pos_teste" in self.dados.columns) and pd.notna(pos) and ("Semana2" in self.dados.columns) and pd.notna(s2):
                s2f = float(s2)
                posf = float(pos)
                retf = float(ret)
                vals = [pre, s1, s2f, posf]
                msk = self._mask_from_present([True, True, True, True])
                X_list.append(vals + msk)
                y_list.append([retf])
                task_list.append("Retencao")
                g_list.append(grupo)
                aidx_list.append(i)

        if not X_list:
            return None, None, None, None, None

        X = np.asarray(X_list, dtype=np.float32)
        y = np.asarray(y_list, dtype=np.float32)
        return X, y, task_list, g_list, aidx_list

    def _normalize_values_in_X(self, X):
        """
        Normaliza SOMENTE as 4 primeiras colunas (valores) usando self.x_mean/self.x_std.
        As 4 colunas finais (máscara) não são normalizadas.
        """
        if (not self.usar_normalizacao) or (self.x_mean is None) or (self.x_std is None):
            return X
        Xn = X.copy()
        Xn[:, :4] = (Xn[:, :4] - self.x_mean) / self.x_std
        return Xn


    # ========================
    # Normalização do alvo por tarefa
    # ========================
    @staticmethod
    def _compute_y_stats_by_task(y, task_list):
        ensure_data_libs()
        stats = {}
        if task_list is None:
            return stats
        y = np.asarray(y, dtype=np.float32).ravel()
        for t in sorted(set(task_list)):
            idx = [i for i, tt in enumerate(task_list) if tt == t]
            if not idx:
                continue
            yt = y[idx]
            mu = float(np.mean(yt))
            sd = float(np.std(yt))
            if sd == 0.0:
                sd = 1.0
            stats[t] = {"mean": mu, "std": sd}
        return stats

    @staticmethod
    def _normalize_y_by_task(y, task_list, stats):
        ensure_data_libs()
        if (task_list is None) or (not stats):
            return np.asarray(y, dtype=np.float32)
        y = np.asarray(y, dtype=np.float32).ravel()
        yn = np.zeros_like(y, dtype=np.float32)
        for i, t in enumerate(task_list):
            st = stats.get(t, None)
            if st is None:
                yn[i] = y[i]
            else:
                yn[i] = (y[i] - float(st["mean"])) / float(st["std"])
        return yn.reshape(-1, 1).astype(np.float32)

    @staticmethod
    def _denormalize_y_by_task(y_norm, task_list, stats):
        ensure_data_libs()
        if (task_list is None) or (not stats):
            return np.asarray(y_norm, dtype=np.float32)
        y_norm = np.asarray(y_norm, dtype=np.float32).ravel()
        y = np.zeros_like(y_norm, dtype=np.float32)
        for i, t in enumerate(task_list):
            st = stats.get(t, None)
            if st is None:
                y[i] = y_norm[i]
            else:
                y[i] = y_norm[i] * float(st["std"]) + float(st["mean"])
        return y.reshape(-1, 1).astype(np.float32)

    def _denorm_single(self, y_norm, task):
        """Desnormaliza um escalar conforme a tarefa.

        Regra:
        - Se existir stats da própria tarefa -> usa.
        - Se NÃO existir (caso comum quando o dataset de treino não tinha aquele alvo),
          faz fallback para stats de uma tarefa compatível para NÃO retornar z-score no relatório.
          Fallback preferencial:
              Retencao -> Pos_teste -> Semana2 -> (qualquer disponível)
        """
        try:
            # sem dicionário => nada a fazer
            if not isinstance(self.y_stats, dict) or not self.y_stats:
                return float(y_norm)

            st = self.y_stats.get(task, None)

            # Fallback de tarefa (evita Ret(Prev) em escala normalizada, tipo -1.0 / 0.2 / 1.5)
            if st is None:
                if task == "Retencao":
                    st = self.y_stats.get("Pos_teste") or self.y_stats.get("Semana2")
                elif task == "Pos_teste":
                    st = self.y_stats.get("Semana2")
                if st is None:
                    # pega qualquer stats disponível
                    try:
                        st = next(iter(self.y_stats.values()))
                    except Exception:
                        st = None

            if st is None:
                return float(y_norm)

            y = float(y_norm) * float(st.get("std", 1.0)) + float(st.get("mean", 0.0))

            # trava em faixa de notas (evita números absurdos no relatório/Word)
            if np.isfinite(y):
                y = float(np.clip(y, 0.0, 100.0))
            return float(y)
        except Exception:
            try:
                return float(np.clip(float(y_norm), 0.0, 100.0))
            except Exception:
                return float(y_norm)


    # ========================
    # Treino
    # ========================
    def _make_split(self):
        self.idx_train, self.idx_val = train_val_split_stratified(
            self.dados, label_col="Grupo", test_size=self.test_size, seed=self.seed
        )

    def _preparar_Xy_com_split_invariavel(self):
        """
        Split por aluno -> constrói amostras (multitarefa) dentro de cada partição.
        Normalização calculada somente com valores das amostras de TREINO (colunas 0..3), ignorando máscara.
        """
        ensure_data_libs()

        self._make_split()

        built_tr = self._build_samples_for_indices(self.idx_train)
        built_va = self._build_samples_for_indices(self.idx_val)

        if built_tr[0] is None:
            raise RuntimeError("Não há amostras de TREINO suficientes (verifique Pre_teste/Semana1 e alvos).")
        if built_va[0] is None:
            raise RuntimeError("Não há amostras de VALIDAÇÃO suficientes (split gerou partição sem alvos). Tente outro seed.")

        X_tr, y_tr, task_tr, g_tr, aidx_tr = built_tr
        X_va, y_va, task_va, g_va, aidx_va = built_va

        if self.usar_normalizacao:
            # calcula mean/std só nos 4 valores (não na máscara)
            vals_tr = X_tr[:, :4]
            self.x_mean = vals_tr.mean(axis=0).astype(np.float32)
            self.x_std = vals_tr.std(axis=0).astype(np.float32)
            self.x_std[self.x_std == 0] = 1.0

            X_tr = self._normalize_values_in_X(X_tr)
            X_va = self._normalize_values_in_X(X_va)
        else:
            self.x_mean = None
            self.x_std = None

        return (X_tr, y_tr, task_tr, g_tr, aidx_tr), (X_va, y_va, task_va, g_va, aidx_va)

    def treinar(self):
        if self.dados is None:
            messagebox.showwarning("Aviso", "Carregue os dados.")
            return
        if not self._tem_dados_para_treino_invariavel():
            messagebox.showwarning(
                "Aviso",
                "Treino exige: Pre_teste + Semana1 e pelo menos 1 alvo preenchido (Semana2 ou Pós ou Retenção)."
            )
            return

        ensure_torch_libs()
        ensure_data_libs()

        self.btn_treinar.config(state="disabled")
        self.status_var.set("Treinando modelo único (multialvo via máscara) com validação e early stopping...")
        set_seed(self.seed)

        try:
            (X_tr, y_tr, task_tr, g_tr, aidx_tr), (X_va, y_va, task_va, g_va, aidx_va) = self._preparar_Xy_com_split_invariavel()

            # Stats do alvo por tarefa (calculadas SOMENTE no treino)
            self.y_stats = self._compute_y_stats_by_task(y_tr, task_tr)

            # Normaliza y por tarefa (evita colapso para a média quando misturamos tarefas)
            y_tr_n = self._normalize_y_by_task(y_tr, task_tr, self.y_stats)
            y_va_n = self._normalize_y_by_task(y_va, task_va, self.y_stats)

            Xtr_t = torch.from_numpy(X_tr)
            ytr_t = torch.from_numpy(y_tr_n)
            Xva_t = torch.from_numpy(X_va)
            yva_t = torch.from_numpy(y_va_n)

            # Pesos por amostra (balanceamento por tarefa)
            counts = {}
            for t in task_tr:
                counts[t] = counts.get(t, 0) + 1
            w = np.array([1.0 / float(counts.get(t, 1)) for t in task_tr], dtype=np.float32)
            w = w / (np.mean(w) if np.mean(w) != 0 else 1.0)  # normaliza (média=1)
            wtr_t = torch.from_numpy(w.reshape(-1, 1))

            Model = get_model_class()
            self.modelo = Model()
            otimizador = optim.Adam(self.modelo.parameters(), lr=0.01)

            best_val = float("inf")
            best_state = None
            patience_cnt = 0
            max_epochs = 5000

            for _ in range(1, max_epochs + 1):
                self.modelo.train()
                saida_tr = self.modelo(Xtr_t)
                perda_tr = torch.mean(wtr_t * (saida_tr - ytr_t) ** 2)
                otimizador.zero_grad()
                perda_tr.backward()
                otimizador.step()

                self.modelo.eval()
                with torch.no_grad():
                    saida_va = self.modelo(Xva_t)
                    perda_va = torch.mean((saida_va - yva_t) ** 2).item()

                if best_val - perda_va > self.min_delta:
                    best_val = perda_va
                    best_state = {k: v.clone() for k, v in self.modelo.state_dict().items()}
                    patience_cnt = 0
                else:
                    patience_cnt += 1

                if patience_cnt >= self.patience:
                    break

            if best_state is not None:
                self.modelo.load_state_dict(best_state)

            # Predições e métricas
            self.modelo.eval()
            with torch.no_grad():
                ytr_pred_n = self.modelo(Xtr_t).numpy().ravel()
                yva_pred_n = self.modelo(Xva_t).numpy().ravel()

            # Desnormaliza por tarefa para métricas no espaço original (0..100)
            ytr = y_tr.ravel()
            yva = y_va.ravel()
            ytr_pred = self._denormalize_y_by_task(ytr_pred_n, task_tr, self.y_stats).ravel()
            yva_pred = self._denormalize_y_by_task(yva_pred_n, task_va, self.y_stats).ravel()

            mse_tr = float(np.mean((ytr - ytr_pred) ** 2))
            r2_tr = r2_score_np(ytr, ytr_pred)
            mse_va = float(np.mean((yva - yva_pred) ** 2))
            r2_va = r2_score_np(yva, yva_pred)

            # Calibração baseada no TREINO (usa as amostras de treino, por grupo)
            self.ajustar_calibracao_com_base_amostras(X_tr, y_tr, g_tr, task_tr)

            # Métricas por tarefa
            by_task_tr = self._metricas_por_tarefa(ytr, ytr_pred, task_tr)
            by_task_va = self._metricas_por_tarefa(yva, yva_pred, task_va)

            self.cache["train"].update({"y_true": ytr, "y_pred": ytr_pred, "mse": mse_tr, "r2": r2_tr, "task": task_tr, "by_task": by_task_tr})
            self.cache["val"].update({"y_true": yva, "y_pred": yva_pred, "mse": mse_va, "r2": r2_va, "task": task_va, "by_task": by_task_va})

            self.view_subset_var.set("val")
            self.view_subset = "val"
            self.update_view_data_and_plot()

            self.status_var.set(
                f"✅ Treinado! MSE_tr={mse_tr:.2f} | R²_tr={r2_tr:.2f} | MSE_val={mse_va:.2f} | R²_val={r2_va:.2f}"
                + "\n" + self._format_by_task_summary(by_task_va, prefix="Val")
            )
            self.btn_salvar.config(state="normal")
            self.btn_atualizar.config(state="normal")
            self.habilitar_previsao_individual()
            self.gerar_relatorio()

        except Exception as e:
            messagebox.showerror("Erro", f"Falha no treinamento:\n{str(e)}")
            self.status_var.set("❌ Erro no treinamento")
        finally:
            pode_treinar = self._tem_dados_para_treino_invariavel()
            self.btn_treinar.config(state="normal" if pode_treinar else "disabled")


    @staticmethod
    def _format_by_task_summary(by_task, prefix=""):
        """Formata métricas por tarefa para exibição em texto."""
        if not by_task:
            return f"{prefix}: sem métricas por tarefa"
        order = ["Semana2", "Pos_teste", "Retencao"]
        parts = []
        for t in order:
            if t in by_task:
                d = by_task[t]
                parts.append(f"{t}: n={d.get('n',0)} R²={d.get('r2',0):.2f} MSE={d.get('mse',0):.2f}")
        # inclui outras tarefas (se houver)
        for t, d in by_task.items():
            if t not in order:
                parts.append(f"{t}: n={d.get('n',0)} R²={d.get('r2',0):.2f} MSE={d.get('mse',0):.2f}")
        return (prefix + " | " if prefix else "") + " ; ".join(parts)

    @staticmethod
    def _metricas_por_tarefa(y_true, y_pred, task_list):
        ensure_data_libs()
        out = {}
        if task_list is None:
            return out
        tasks = sorted(set(task_list))
        for t in tasks:
            idx = [i for i, tt in enumerate(task_list) if tt == t]
            if not idx:
                continue
            yt = np.asarray([y_true[i] for i in idx], dtype=float)
            yp = np.asarray([y_pred[i] for i in idx], dtype=float)
            mse = float(np.mean((yt - yp) ** 2))
            r2 = r2_score_np(yt, yp)
            out[t] = {"n": int(len(idx)), "mse": mse, "r2": r2}
        return out

    def ajustar_calibracao_com_base_amostras(self, X_tr, y_tr, grupos_tr, task_tr=None):
        """Calibração usando as amostras de treino já montadas (evita reconstrução)."""
        ensure_torch_libs()
        ensure_data_libs()

        if self.modelo is None:
            return

        X_use = X_tr
        if self.usar_normalizacao and (self.x_mean is not None):
            X_use = self._normalize_values_in_X(X_use)

        with torch.no_grad():
            y_hat = self.modelo(torch.from_numpy(X_use)).numpy().ravel()

        # Se o modelo estiver treinando em y normalizado por tarefa, desnormaliza aqui
        if task_tr is not None and isinstance(task_tr, (list, tuple)) and self.y_stats:
            y_hat = self._denormalize_y_by_task(y_hat, task_tr, self.y_stats).ravel()

        y_true = y_tr.ravel()
        y_pred = y_hat.ravel()
        grupos = np.asarray(grupos_tr, dtype=object)

        my, myh = float(np.mean(y_true)), float(np.mean(y_pred))
        self.calib["global"] = (my / myh) if myh != 0 else 1.0

        self.calib["group_alpha"].clear()
        self.calib["group_lin"].clear()
        for g in np.unique(grupos):
            idx = np.where(grupos == g)[0]
            y_g = y_true[idx]
            yhat_g = y_pred[idx]
            myg, myhg = float(np.mean(y_g)), float(np.mean(yhat_g))
            self.calib["group_alpha"][g] = (myg / myhg) if myhg != 0 else 1.0

            var = float(np.var(yhat_g))
            if var > 0:
                cov = float(np.mean((yhat_g - np.mean(yhat_g)) * (y_g - np.mean(y_g))))
                a = cov / var
                b = float(np.mean(y_g)) - a * float(np.mean(yhat_g))
            else:
                a, b = 1.0, 0.0
            self.calib["group_lin"][g] = (a, b)

    def atualizar_modelo(self):
        self.treinar()

    # ========================
    # Visual resultados
    # ========================
    def mostrar_resultados_modelo(self, mse, r2, titulo_subset="(Validação)"):
        ensure_plot_libs()
        ensure_data_libs()
        plt.close("all")

        info = "Modelo: Rede Neural (8 → 32 → 16 → 1)\n"
        info += "Entradas: [Pre, S1, S2, Pós] + máscara binária (4)\n"
        info += f"Split: Treino/Val = {int((1 - self.test_size) * 100)}/{int(self.test_size * 100)} | Seed={self.seed}\n"
        info += f"Early stopping: patience={self.patience}, min_delta={self.min_delta}\n"
        info += f"Normalização (apenas valores): {'Sim' if self.usar_normalizacao else 'Não'}\n"
        info += f"MSE {titulo_subset}: {mse:.2f} | R² {titulo_subset}: {r2:.2f}\n"
        info += f"Calibração: {self.calib_method}\n\n"
        info += "Métricas por tarefa (n, MSE, R²):\n"
        part = "val" if "Validação" in titulo_subset else "train"
        bt = self.cache.get(part, {}).get("by_task", {})
        if bt:
            for t in ["Semana2", "Pos_teste", "Retencao"]:
                if t in bt:
                    info += f"• {t:<9} n={bt[t]['n']:<4d} MSE={bt[t]['mse']:.2f} R²={bt[t]['r2']:.2f}\n"
        else:
            info += "• (indisponível)\n"

        self.info_modelo.config(state=tk.NORMAL)
        self.info_modelo.delete(1.0, tk.END)
        self.info_modelo.insert(tk.END, info)
        self.info_modelo.config(state=tk.DISABLED)

        # Gráfico geral: Real vs Previsto (mistura tarefas, mas útil como sanity-check)
        fig2, (ax1, ax2) = plt.subplots(1, 2, figsize=(12, 5))
        ax1.plot(self.y_real, label="Real", marker="o")
        ax1.plot(self.y_previsto, label="Previsto (cru)", marker="x", linestyle="--")
        ax1.set_title(f"Real vs Previsto {titulo_subset} (todas as tarefas)")
        ax1.legend()
        ax1.set_xlabel("Amostra")
        ax1.set_ylabel("Nota")

        lo = float(min(np.min(self.y_real), np.min(self.y_previsto))) - 2
        hi = float(max(np.max(self.y_real), np.max(self.y_previsto))) + 2
        ax2.scatter(self.y_real, self.y_previsto, alpha=0.7)
        ax2.plot([lo, hi], [lo, hi], "r--", lw=2)
        ax2.set_xlim(lo, hi)
        ax2.set_ylim(lo, hi)
        ax2.set_xlabel("Real")
        ax2.set_ylabel("Previsto (cru)")
        ax2.set_title(f"R² {titulo_subset} = {r2:.2f}")

        fig2.suptitle("Resultados do Modelo Único", fontsize=16)
        fig2.tight_layout()

        for widget in self.frame_grafico2.winfo_children():
            widget.destroy()

        canvas2 = FigureCanvasTkAgg(fig2, self.frame_grafico2)
        canvas2.draw()
        canvas2.get_tk_widget().pack(pady=10)
        self.canvas2 = canvas2

    # =======================
    # Salvar/Carregar modelo
    # =======================
    def salvar_modelo(self):
        ensure_torch_libs()
        ensure_data_libs()
        if not self.modelo:
            messagebox.showwarning("Aviso", "Nenhum modelo para salvar.")
            return

        def _to_py(obj):
            """Converte numpy/tensors para tipos Python seguros (compatível com torch>=2.6)."""
            try:
                if obj is None:
                    return None
                if hasattr(obj, "detach") and hasattr(obj, "cpu") and hasattr(obj, "numpy"):
                    obj = obj.detach().cpu().numpy()
                if hasattr(obj, "shape") and hasattr(obj, "dtype") and hasattr(obj, "tolist"):
                    return obj.tolist()
                if hasattr(obj, "item") and not isinstance(obj, (dict, list, tuple, str, bytes)):
                    try:
                        return obj.item()
                    except Exception:
                        pass
                if isinstance(obj, dict):
                    return {str(k): _to_py(v) for k, v in obj.items()}
                if isinstance(obj, list):
                    return [_to_py(v) for v in obj]
                if isinstance(obj, tuple):
                    return tuple(_to_py(v) for v in obj)
                return obj
            except Exception:
                return obj

        payload = {
            "state_dict": self.modelo.state_dict(),
            "x_mean": _to_py(self.x_mean),
            "x_std": _to_py(self.x_std),
            "usar_normalizacao": bool(self.usar_normalizacao),
            "calib": _to_py(self.calib),
            "calib_method": str(self.calib_method),
            "y_stats": _to_py(self.y_stats),
        }

        torch.save(payload, self.caminho_modelo)
        messagebox.showinfo("Sucesso", f"Modelo salvo como:\n{self.caminho_modelo}")
    def carregar_modelo(self):
        caminho = filedialog.askopenfilename(title="Selecione o modelo (.pth)", filetypes=[("PyTorch Model", "*.pth")])
        if not caminho:
            return
        try:
            ensure_torch_libs()
            Model = get_model_class()
            self.modelo = Model()

            # PyTorch 2.6+ usa weights_only=True por padrão; checkpoints antigos (com numpy) podem falhar.
            # Estratégia: tenta modo seguro (weights_only=True). Se falhar e o arquivo for confiável, cai para weights_only=False.
            try:
                payload = torch.load(caminho, map_location=torch.device('cpu'), weights_only=True)
            except TypeError:
                # versões antigas do torch não têm weights_only
                payload = torch.load(caminho, map_location=torch.device('cpu'))
            except Exception as e_safe:
                try:
                    payload = torch.load(caminho, map_location=torch.device('cpu'), weights_only=False)
                    messagebox.showwarning(
                        'Aviso',
                        'Checkpoint carregado em modo compatível (weights_only=False).\nUse apenas modelos gerados por você/por fonte confiável.'
                    )
                except Exception as e_full:
                    raise RuntimeError(f'Falha ao carregar checkpoint (safe): {e_safe}\\nFalha (full): {e_full}')
            if isinstance(payload, dict) and "state_dict" in payload:
                self.modelo.load_state_dict(payload["state_dict"])
                self.x_mean = payload.get("x_mean", None)
                self.x_std = payload.get("x_std", None)
                # converte listas salvas para numpy
                try:
                    if isinstance(self.x_mean, list):
                        self.x_mean = np.asarray(self.x_mean, dtype=np.float32)
                    if isinstance(self.x_std, list):
                        self.x_std = np.asarray(self.x_std, dtype=np.float32)
                except Exception:
                    pass
                self.usar_normalizacao = bool(payload.get("usar_normalizacao", True))
                self.calib = payload.get("calib", self.calib)
                self.calib_method = payload.get("calib_method", self.calib_method)
                self.y_stats = payload.get("y_stats", {})
            else:
                # fallback: arquivo antigo com apenas state_dict
                self.modelo.load_state_dict(payload)

            self.modelo.eval()
            self.caminho_modelo = caminho
            self.status_var.set("✅ Modelo carregado com sucesso!")
            self.btn_salvar.config(state="normal")

            pode_treinar = self._tem_dados_para_treino_invariavel()
            self.btn_treinar.config(state="normal" if pode_treinar else "disabled")
            self.btn_atualizar.config(state="normal" if pode_treinar else "disabled")

            self.habilitar_previsao_individual()

            self.cache = {
                "train": {"y_true": None, "y_pred": None, "mse": None, "r2": None, "task": None, "by_task": {}},
                "val": {"y_true": None, "y_pred": None, "mse": None, "r2": None, "task": None, "by_task": {}},
            }
            self.view_subset_var.set("val")
            self.view_subset = "val"

            messagebox.showinfo("Sucesso", "Modelo carregado! Para métricas/validação, treine com seu CSV.")
        except Exception as e:
            messagebox.showerror("Erro", f"Falha ao carregar modelo:\n{str(e)}")

    # =======================
    # Previsão individual
    # =======================
    def aplicar_calibracao(self, yhat, grupo):
        if self.calib_method == "alpha_global":
            alpha = self.calib.get("global", 1.0)
            return alpha * yhat
        elif self.calib_method == "alpha_group":
            alpha = self.calib["group_alpha"].get(grupo, 1.0)
            return alpha * yhat
        elif self.calib_method == "lin_group":
            a, b = self.calib["group_lin"].get(grupo, (1.0, 0.0))
            return a * yhat + b
        else:
            return yhat

    def _infer_task_from_inputs(self, pre, s1, s2, pos):
        """
        Regra de decisão (invariável) baseada no que foi informado:
          - pre+s1 -> Semana2
          - pre+s1+s2 -> Pos_teste
          - pre+s1+s2+pos -> Retencao
        """
        has_pre = pre is not None
        has_s1 = s1 is not None
        has_s2 = s2 is not None
        has_pos = pos is not None
        if not (has_pre and has_s1):
            return None
        if has_pre and has_s1 and (not has_s2) and (not has_pos):
            return "Semana2"
        if has_pre and has_s1 and has_s2 and (not has_pos):
            return "Pos_teste"
        if has_pre and has_s1 and has_s2 and has_pos:
            return "Retencao"
        return None

    def prever_individual(self):
        try:
            if self.modelo is None:
                raise RuntimeError("Carregue um modelo primeiro.")

            ensure_torch_libs()
            ensure_data_libs()

            grupo = self.var_grupo.get()

            def _get_float_or_none(s):
                s = (s or "").strip()
                if s == "":
                    return None
                return float(s)

            pre = _get_float_or_none(self.var_pre.get())
            sem1 = _get_float_or_none(self.var_sem1.get())
            sem2 = _get_float_or_none(self.var_sem2.get())
            pos = _get_float_or_none(self.var_pos.get())

            # checa range apenas para os preenchidos
            for v in [pre, sem1, sem2, pos]:
                if v is None:
                    continue
                if not (0 <= v <= 100):
                    raise ValueError("Notas devem estar entre 0 e 100")

            task = self._infer_task_from_inputs(pre, sem1, sem2, pos)
            if task is None:
                raise ValueError(
                    "Sequência inválida. Use uma das opções:\n"
                    "1) Pre + Semana1\n"
                    "2) Pre + Semana1 + Semana2\n"
                    "3) Pre + Semana1 + Semana2 + Pós"
                )

            # monta X + máscara
            v_pre = float(pre)
            v_s1 = float(sem1)
            v_s2 = float(sem2) if sem2 is not None else 0.0
            v_pos = float(pos) if pos is not None else 0.0

            present = [True, True, sem2 is not None, pos is not None]
            msk = self._mask_from_present(present)
            X = np.array([[v_pre, v_s1, v_s2, v_pos] + msk], dtype=np.float32)

            # normaliza valores
            X = self._normalize_values_in_X(X)

            with torch.no_grad():
                yhat_n = self.modelo(torch.from_numpy(X)).item()
            yhat_cru = self._denorm_single(yhat_n, task)

            yhat = self.aplicar_calibracao(yhat_cru, grupo)

            label = {"Semana2": "Semana 2", "Pos_teste": "Pós-teste", "Retencao": "Retenção"}[task]
            self.resultado_var.set(f"🎯 {label} prevista: {yhat:.1f}")
            self.dica_label.config(foreground="green", text=f"Alvo detectado: {label} | Grupo: {grupo}")

        except Exception as e:
            messagebox.showerror("Erro", str(e))

    def habilitar_previsao_individual(self):
        if self.modelo:
            self.btn_prever.config(state="normal")
            self.dica_label.config(text="✅ Modelo carregado! Pronto para prever.", foreground="green")
        else:
            self.btn_prever.config(state="disabled")
            self.dica_label.config(text="💡 Carregue um modelo para habilitar a previsão", foreground="red")

    # =======================
    # Exportações
    # =======================
    def _predict_one_task(self, pre, s1, s2, pos, grupo, task):
        """Prediz um único task, dado valores (podem ser nan), usando máscara coerente."""
        ensure_data_libs()
        ensure_torch_libs()

        if self.modelo is None:
            return np.nan

        if task == "Semana2":
            if pd.isna(pre) or pd.isna(s1):
                return np.nan
            vals = [float(pre), float(s1), 0.0, 0.0]
            msk = self._mask_from_present([True, True, False, False])
        elif task == "Pos_teste":
            if pd.isna(pre) or pd.isna(s1) or pd.isna(s2):
                return np.nan
            vals = [float(pre), float(s1), float(s2), 0.0]
            msk = self._mask_from_present([True, True, True, False])
        elif task == "Retencao":
            if pd.isna(pre) or pd.isna(s1) or pd.isna(s2) or pd.isna(pos):
                return np.nan
            vals = [float(pre), float(s1), float(s2), float(pos)]
            msk = self._mask_from_present([True, True, True, True])
        else:
            return np.nan

        X = np.array([vals + msk], dtype=np.float32)
        X = self._normalize_values_in_X(X)
        with torch.no_grad():
            yhat_n = self.modelo(torch.from_numpy(X)).item()
        yhat_cru = self._denorm_single(yhat_n, task)
        yhat_cal = float(self.aplicar_calibracao(yhat_cru, grupo))
        # trava faixa de nota
        if np.isfinite(yhat_cal):
            yhat_cal = float(np.clip(yhat_cal, 0.0, 100.0))
        return float(yhat_cal)

    def exportar_excel(self):
        if self.dados is None or self.modelo is None:
            messagebox.showwarning("Aviso", "Carregue dados e modelo.")
            return

        ensure_data_libs()

        df_export = self.dados.copy()

        # Gera previsões coerentes para cada tarefa
        preds_s2 = []
        preds_pos = []
        preds_ret = []

        for _, row in df_export.iterrows():
            grupo = str(row.get("Grupo", ""))
            pre = row.get("Pre_teste", np.nan)
            s1 = row.get("Semana1", np.nan)
            s2 = row.get("Semana2", np.nan)
            pos = row.get("Pos_teste", np.nan)

            preds_s2.append(self._predict_one_task(pre, s1, s2, pos, grupo, "Semana2"))
            preds_pos.append(self._predict_one_task(pre, s1, s2, pos, grupo, "Pos_teste"))
            preds_ret.append(self._predict_one_task(pre, s1, s2, pos, grupo, "Retencao"))

        df_export["Prev_Semana2"] = preds_s2
        df_export["Prev_Pos"] = preds_pos
        df_export["Prev_Retencao"] = preds_ret

        caminho = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
        if caminho:
            df_export.to_excel(caminho, index=False)
            messagebox.showinfo("Sucesso", f"Planilha salva como:\n{caminho}")

    def exportar_word(self):
        try:
            ensure_data_libs()
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
        except ImportError:
            messagebox.showerror("Erro", "Instale python-docx: pip install python-docx")
            return

        caminho = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if not caminho:
            return

        def _set_cell_shading(cell, fill_hex: str):
            # fill_hex: e.g. "D9E1F2"
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill_hex)
            tcPr.append(shd)

        def _set_run_font(run, name="Times New Roman", size_pt=12, bold=None, italic=None):
            run.font.name = name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
            run.font.size = Pt(size_pt)
            if bold is not None:
                run.bold = bold
            if italic is not None:
                run.italic = italic

        def _add_h1(texto):
            p = doc.add_paragraph()
            run = p.add_run(texto)
            _set_run_font(run, size_pt=14, bold=True)
            p.space_before = Pt(10)
            p.space_after = Pt(4)
            return p

        def _add_h2(texto):
            p = doc.add_paragraph()
            run = p.add_run(texto)
            _set_run_font(run, size_pt=12, bold=True)
            p.space_before = Pt(8)
            p.space_after = Pt(2)
            return p

        def _add_body(texto):
            p = doc.add_paragraph()
            run = p.add_run(texto)
            _set_run_font(run, size_pt=12)
            p.space_after = Pt(1)
            return p

        def _add_table(rows, header=True, col_align=None):
            # rows: List[List[str]]
            if not rows:
                return None
            ncols = max(len(r) for r in rows)
            table = doc.add_table(rows=1 if header else 0, cols=ncols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"

            def _write_row(row_cells, data, is_header=False):
                for j in range(ncols):
                    txt = "" if j >= len(data) else str(data[j])
                    cell = row_cells[j]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(txt)
                    _set_run_font(run, size_pt=11, bold=is_header)
                    if col_align and j < len(col_align):
                        p.alignment = col_align[j]
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    if is_header:
                        _set_cell_shading(cell, "D9E1F2")

            start_idx = 0
            if header:
                hdr = rows[0]
                _write_row(table.rows[0].cells, hdr, is_header=True)
                start_idx = 1

            for i in range(start_idx, len(rows)):
                row = table.add_row().cells
                _write_row(row, rows[i], is_header=False)

            doc.add_paragraph()
            return table

        # =========================
        # Documento + layout
        # =========================
        doc = Document()

        # Margens (estilo acadêmico)
        secao = doc.sections[0]
        secao.top_margin = Inches(0.9)
        secao.bottom_margin = Inches(0.9)
        secao.left_margin = Inches(1.0)
        secao.right_margin = Inches(1.0)

        # Fonte padrão
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")
        style.font.size = Pt(12)

        # =========================
        # Capa
        # =========================
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_title.add_run("Relatório de Desempenho Educacional")
        _set_run_font(run, size_pt=16, bold=True)

        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_date.add_run(f"Data: {pd.Timestamp.now().strftime('%d/%m/%Y')}")
        _set_run_font(run, size_pt=11, italic=True)

        doc.add_paragraph()

        # =========================
        # Parser do texto bruto
        # =========================
        raw = self.txt_relatorio.get(1.0, tk.END).strip("\n")
        linhas = [l.rstrip() for l in raw.split("\n")]

        def is_sep(l):
            s = l.strip()
            return (set(s) <= set("=—-") and len(s) >= 10)

        # Remover linhas vazias extras e separadores repetidos
        clean = []
        for l in linhas:
            if not l.strip():
                continue
            if is_sep(l):
                # ignora separadores do modo "console"
                continue
            clean.append(l)

        # Estado do parser
        i = 0
        while i < len(clean):
            line = clean[i].strip()

            # Títulos de seção
            if line.upper().startswith("RELATÓRIO DE DESEMPENHO"):
                _add_h1("RELATÓRIO DE DESEMPENHO")
                i += 1
                continue

            if line.startswith("COLUNAS DISPONÍVEIS"):
                _add_h2("Colunas disponíveis para curva (dataset)")
                i += 1
                # próximos itens (bullets) até encontrar próximo cabeçalho conhecido
                while i < len(clean) and not clean[i].strip().upper().startswith(("TOP 7", "ERROS POR", "MÉTRICAS", "🔍", "⚠️")):
                    _add_body(clean[i].lstrip("• ").strip())
                    i += 1
                doc.add_paragraph()
                continue

            if line.startswith("TOP 7 TEMAS"):
                _add_h2("Top 7 temas com mais erros (questões)")
                i += 1
                # Se indisponível, imprime o texto e segue
                if i < len(clean) and "Indisponível" in clean[i]:
                    _add_body(clean[i])
                    i += 1
                    doc.add_paragraph()
                    continue

                # Tabela: Tema | Erros | Total | %Erro
                # Espera linha de cabeçalho na sequência
                rows = []
                while i < len(clean) and clean[i].strip() and not clean[i].strip().upper().startswith(("ERROS POR", "MÉTRICAS", "🔍", "⚠️")):
                    rows.append(clean[i])
                    i += 1
                # parse simples por múltiplos espaços
                parsed = []
                for r in rows:
                    if "Tema" in r and "Erro" in r:
                        # cabeçalho
                        parts = re.split(r"\s{2,}", r.strip())
                        parsed.append(parts)
                    else:
                        parts = re.split(r"\s{2,}", r.strip())
                        if len(parts) >= 2:
                            parsed.append(parts)
                if len(parsed) >= 2:
                    _add_table(parsed, header=True)
                else:
                    for r in rows:
                        _add_body(r)
                continue

            if line.startswith("ERROS POR NÍVEL COGNITIVO"):
                _add_h2("Erros por nível cognitivo (questões)")
                i += 1
                if i < len(clean) and "Indisponível" in clean[i]:
                    _add_body(clean[i])
                    i += 1
                    doc.add_paragraph()
                    continue

                rows = []
                while i < len(clean) and clean[i].strip() and not clean[i].strip().upper().startswith(("MÉTRICAS", "🔍", "⚠️")):
                    rows.append(clean[i])
                    i += 1

                parsed = []
                for r in rows:
                    parts = re.split(r"\s{2,}", r.strip())
                    if len(parts) >= 2:
                        parsed.append(parts)
                if len(parsed) >= 2:
                    _add_table(parsed, header=True)
                else:
                    for r in rows:
                        _add_body(r)
                continue

            if line.startswith("TEMAS PRIORITÁRIOS"):
                _add_h2("Temas prioritários (maior impacto de erro)")
                i += 1

                rows = []
                while i < len(clean) and clean[i].strip() and not clean[i].strip().upper().startswith(("TENDÊNCIA TEMPORAL", "MÉTRICAS", "🔍", "⚠️")) and not clean[i].strip().startswith("Leitura rápida:"):
                    rows.append(clean[i])

                parsed = []
                for r in rows:
                    if set(r.strip()) == {"-"}:
                        continue
                    parts = re.split(r"\s{2,}", r.strip())
                    if parts and any(p for p in parts):
                        parsed.append(parts)

                if len(parsed) >= 2:
                    _add_table(parsed, header=True)
                else:
                    for r in rows:
                        _add_body(r)

                doc.add_paragraph()
                continue

            if line.startswith("TENDÊNCIA TEMPORAL"):
                _add_h2("Tendência temporal (médias e variações)")
                i += 1

                rows = []
                while i < len(clean) and clean[i].strip() and not clean[i].strip().upper().startswith(("MÉTRICAS", "🔍", "⚠️")) and not clean[i].strip().startswith("Leitura rápida:"):
                    rows.append(clean[i])

                parsed = []
                for r in rows:
                    if set(r.strip()) == {"-"}:
                        continue
                    parts = re.split(r"\s{2,}", r.strip())
                    if parts and any(p for p in parts):
                        parsed.append(parts)

                if len(parsed) >= 2:
                    _add_table(parsed, header=True)
                else:
                    for r in rows:
                        _add_body(r)

                # Se existir leitura rápida logo abaixo, imprime
                if i < len(clean) and clean[i].strip().startswith("Leitura rápida:"):
                    _add_body(clean[i].strip())
                    i += 1

                doc.add_paragraph()
                continue

                # Se existir leitura rápida logo abaixo, imprime
                if i < len(clean) and clean[i].strip().startswith("Leitura rápida:"):
                    _add_body(clean[i].strip())
                    i += 1

                doc.add_paragraph()
                continue
            if line.startswith("MÉTRICAS DO MODELO (GLOBAL)"):
                _add_h2("Métricas do modelo (global)")
                i += 1
                rows = [["Conjunto", "MSE", "R²"]]
                while i < len(clean) and clean[i].strip().startswith(("Treino", "Validação", "Validacao", "Teste")):
                    # Ex: Treino    -> MSE: 2.61 | R²: 0.98
                    s = clean[i]
                    conjunto = s.split("->")[0].strip()
                    mse = ""
                    r2 = ""
                    m1 = re.search(r"MSE:\s*([0-9\.\-]+)", s)
                    m2 = re.search(r"R²:\s*([0-9\.\-]+)", s)
                    if m1: mse = m1.group(1)
                    if m2: r2 = m2.group(1)
                    rows.append([conjunto, mse, r2])
                    i += 1
                _add_table(rows, header=True)
                continue

            if line.startswith("MÉTRICAS POR TAREFA (TREINO)"):
                _add_h2("Métricas por tarefa (treino)")
                i += 1
                rows = [["Tarefa", "n", "MSE", "R²"]]
                while i < len(clean) and clean[i].strip().startswith("•"):
                    s = clean[i].lstrip("•").strip()
                    # Ex: Semana2   n=24   MSE=1.70 | R²=0.96
                    tarefa = s.split("n=")[0].strip()
                    n = ""
                    mse = ""
                    r2 = ""
                    mN = re.search(r"n\s*=\s*([0-9]+)", s)
                    mM = re.search(r"MSE\s*=\s*([0-9\.\-]+)", s)
                    mR = re.search(r"R²\s*=\s*([0-9\.\-]+)", s)
                    if mN: n = mN.group(1)
                    if mM: mse = mM.group(1)
                    if mR: r2 = mR.group(1)
                    rows.append([tarefa, n, mse, r2])
                    i += 1
                _add_table(rows, header=True)
                continue

            if line.startswith("MÉTRICAS POR TAREFA (VALIDAÇÃO)") or line.startswith("MÉTRICAS POR TAREFA (VALIDACAO)"):
                _add_h2("Métricas por tarefa (validação)")
                i += 1
                rows = [["Tarefa", "n", "MSE", "R²"]]
                while i < len(clean) and clean[i].strip().startswith("•"):
                    s = clean[i].lstrip("•").strip()
                    tarefa = s.split("n=")[0].strip()
                    n = ""
                    mse = ""
                    r2 = ""
                    mN = re.search(r"n\s*=\s*([0-9]+)", s)
                    mM = re.search(r"MSE\s*=\s*([0-9\.\-]+)", s)
                    mR = re.search(r"R²\s*=\s*([0-9\.\-]+)", s)
                    if mN: n = mN.group(1)
                    if mM: mse = mM.group(1)
                    if mR: r2 = mR.group(1)
                    rows.append([tarefa, n, mse, r2])
                    i += 1
                _add_table(rows, header=True)
                continue

            if line.startswith("🔍 PREVISÕES"):
                _add_h1("PREVISÕES (MODELO ÚNICO)")
                i += 1
                # Cabeçalho da tabela (linha com colunas)
                # Procurar a primeira linha que tenha "Aluno" e "Grupo"
                while i < len(clean) and not ("Aluno" in clean[i] and "Grupo" in clean[i]):
                    i += 1
                if i >= len(clean):
                    continue
                header_line = clean[i]
                header = header_line.strip().split()
                i += 1

                rows = [header]
                while i < len(clean):
                    l = clean[i].strip()
                    if l.startswith("Legenda:") or l.startswith("⚠️"):
                        break
                    # linha de aluno (pode ter múltiplos espaços)
                    parts = clean[i].strip().split()
                    if len(parts) >= 2:
                        rows.append(parts)
                    i += 1

                # Alinhamento: texto à esquerda; notas à direita
                col_align = []
                for j, h in enumerate(rows[0]):
                    if j <= 1:
                        col_align.append(WD_ALIGN_PARAGRAPH.LEFT)
                    else:
                        col_align.append(WD_ALIGN_PARAGRAPH.RIGHT)

                _add_table(rows, header=True, col_align=col_align)

                # Legenda (se existir)
                while i < len(clean) and clean[i].strip().startswith("Legenda:"):
                    _add_body(clean[i].strip())
                    i += 1
                doc.add_paragraph()
                continue

            if line.startswith("⚠️ RELATÓRIO"):
                _add_h1("ALUNOS COM BAIXO RENDIMENTO (nota < limiar)")
                i += 1
                # Capturar bloco até o fim
                bloco = []
                while i < len(clean):
                    bloco.append(clean[i])
                    i += 1

                # Resumo por grupo (bullets)
                idx = 0
                while idx < len(bloco) and not bloco[idx].startswith("Resumo por Grupo"):
                    idx += 1
                if idx < len(bloco):
                    _add_h2("Resumo por grupo")
                    idx += 1
                    while idx < len(bloco) and bloco[idx].strip().startswith("•"):
                        _add_body(bloco[idx].lstrip("• ").strip())
                        idx += 1

                # Lista detalhada (tabela)
                while idx < len(bloco) and not bloco[idx].startswith("Lista detalhada"):
                    idx += 1
                if idx < len(bloco):
                    _add_h2("Lista detalhada")
                    idx += 1
                    # Cada linha: AlunoX | Grupo: ... | Contribuições: ...
                    rows = [["Aluno", "Grupo", "Contribuições"]]
                    while idx < len(bloco) and not bloco[idx].startswith("Ações pedagógicas sugeridas"):
                        s = bloco[idx].strip()
                        if "|" in s:
                            parts = [p.strip() for p in s.split("|")]
                            aluno = parts[0]
                            grupo = ""
                            contrib = ""
                            for p in parts[1:]:
                                if p.lower().startswith("grupo:"):
                                    grupo = p.split(":", 1)[1].strip()
                                if p.lower().startswith("contribuições:") or p.lower().startswith("contribuicoes:"):
                                    contrib = p.split(":", 1)[1].strip()
                            rows.append([aluno, grupo, contrib])
                        idx += 1

                    # alinhamentos
                    col_align = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT]
                    _add_table(rows, header=True, col_align=col_align)

                # Ações sugeridas
                while idx < len(bloco) and not bloco[idx].startswith("Ações pedagógicas sugeridas"):
                    idx += 1
                if idx < len(bloco):
                    _add_h2("Ações pedagógicas sugeridas")
                    idx += 1
                    while idx < len(bloco):
                        if bloco[idx].strip().startswith("•"):
                            _add_body(bloco[idx].lstrip("• ").strip())
                        idx += 1
                continue

            # Linhas "soltas" (fallback)
            _add_body(line)
            i += 1

        doc.save(caminho)
        messagebox.showinfo("Sucesso", f"Relatório Word salvo como:\n{caminho}")


    def salvar_grafico(self):
        if hasattr(self, "canvas1"):
            caminho = filedialog.asksaveasfilename(defaultextension=".png", filetypes=[("PNG files", "*.png")])
            if caminho:
                self.canvas1.figure.savefig(caminho, dpi=150, bbox_inches="tight")
                messagebox.showinfo("Sucesso", "Gráfico salvo com sucesso!")
        else:
            messagebox.showwarning("Aviso", "Nenhum gráfico para salvar.")

    # =======================
    # Relatórios
    # =======================


    def gerar_relatorio(self):
        if self.dados is None:
            return

        ensure_data_libs()
        self._detectar_colunas_curva()

        df = self.dados.copy()

        relatorio = " " * 20 + "RELATÓRIO DE DESEMPENHO\n"
        relatorio += "=" * 88 + "\n\n"

        # ---------------------------------------------------------------------
        # Colunas disponíveis (curva)
        # ---------------------------------------------------------------------
        relatorio += "COLUNAS DISPONÍVEIS PARA CURVA (DATASET)\n"
        relatorio += "-" * 100 + "\n"
        if self.colunas_disponiveis_curva:
            relatorio += " • " + ", ".join([lab for _, lab in self.colunas_disponiveis_curva]) + "\n\n"
        else:
            relatorio += "Nenhuma coluna válida encontrada.\n\n"

        # ---------------------------------------------------------------------
        # Erros por questões (se existir Q1..Qn)
        # ---------------------------------------------------------------------
        temas = self._calc_erros_agrupado("tema")
        niveis = self._calc_erros_agrupado("nivel")

        # ---------------------------------------------------------------------
        # SUMÁRIO EXECUTIVO (tomada de decisão)
        # ---------------------------------------------------------------------
        relatorio += "SUMÁRIO EXECUTIVO (SUPORTE À DECISÃO)\n"
        relatorio += "-" * 100 + "\n"

        # 1) IFC (Índice de Fragilidade Cognitiva) a partir de níveis
        pesos_nivel = {
            "CPB": 1.0,
            "CPM": 1.2,
            "CAP": 1.4,
        }

        ifc = None
        diag_estrutural = "Indisponível (nível cognitivo não identificado no dataset de questões)."
        if niveis:
            # taxa de erro por nível (0..1)
            taxa_por_nivel = {}
            for k, e, t in niveis:
                kk = str(k).strip().upper()
                taxa_por_nivel[kk] = (float(e) / float(t)) if t else 0.0

            ifc = 0.0
            for kk, pe in taxa_por_nivel.items():
                w = pesos_nivel.get(kk, 1.0)
                ifc += w * pe

            # diagnóstico estrutural simples
            pe_cpb = taxa_por_nivel.get("CPB", None)
            pe_cpm = taxa_por_nivel.get("CPM", None)
            pe_cap = taxa_por_nivel.get("CAP", None)

            if pe_cpb is not None and pe_cpm is not None and pe_cap is not None:
                if pe_cpb >= pe_cpm and pe_cpb >= pe_cap:
                    diag_estrutural = "Predomínio de erro em CPB (base conceitual). Prioridade: reforço de pré-requisitos + prática guiada."
                elif pe_cap >= pe_cpb and pe_cap >= pe_cpm:
                    diag_estrutural = "Predomínio de erro em CAP (aplicação/transferência). Prioridade: problemas contextualizados + projetos/situações reais."
                else:
                    diag_estrutural = "Predomínio intermediário (CPM). Prioridade: mediação, exemplos trabalhados e feedback rápido."

        # 2) Risco do grupo (alunos < 40)
        limiar_baixo = 40.0
        col_pos = "Pos_teste" if "Pos_teste" in df.columns else None
        col_ret = "Retencao" if "Retencao" in df.columns else None

        pct_baixo = None
        n_baixo = None
        n_total = int(len(df))

        if col_pos is not None:
            pos_vals = pd.to_numeric(df[col_pos], errors="coerce")
            n_valid = int(pos_vals.notna().sum())
            if n_valid > 0:
                n_baixo = int((pos_vals < limiar_baixo).sum())
                pct_baixo = float(n_baixo) / float(n_valid)

        # 3) Status geral (ESTÁVEL / ALERTA / CRÍTICO)
        status = "INDETERMINADO"
        motivos = []

        # Regras simples e conservadoras
        if ifc is not None and ifc >= 1.5:
            motivos.append(f"IFC={ifc:.2f} (crítico ≥ 1.50)")
        elif ifc is not None and ifc >= 1.0:
            motivos.append(f"IFC={ifc:.2f} (alerta ≥ 1.00)")

        if pct_baixo is not None:
            motivos.append(f"Alunos abaixo de {int(limiar_baixo)} no Pós: {n_baixo}/{n_valid} ({pct_baixo*100:.1f}%)")

        # decisão
        critico = False
        alerta = False
        if ifc is not None and ifc >= 1.5:
            critico = True
        if pct_baixo is not None and pct_baixo >= 0.40:
            critico = True
        if col_pos is not None:
            mpos = float(pd.to_numeric(df[col_pos], errors="coerce").mean())
            if not math.isnan(mpos) and mpos < 60:
                alerta = True
                motivos.append(f"Média do Pós={mpos:.1f} (<60)")
            if not math.isnan(mpos) and mpos < 50:
                critico = True

        if critico:
            status = "CRÍTICO"
        elif alerta or (ifc is not None and ifc >= 1.0) or (pct_baixo is not None and pct_baixo >= 0.20):
            status = "ALERTA"
        else:
            status = "ESTÁVEL"

        relatorio += f"Status do grupo: {status}\n"
        if ifc is not None:
            relatorio += f"IFC (Índice de Fragilidade Cognitiva): {ifc:.2f}\n"
        relatorio += f"Diagnóstico estrutural: {diag_estrutural}\n"
        if motivos:
            relatorio += "Sinais/justificativas: " + " | ".join(motivos[:4]) + "\n"
        relatorio += "\n"

        
        # 4) Tabelas decisórias (separadas)
        #    (A) Temas prioritários (maior impacto de erro)
        #    (B) Tendência temporal (médias e variações)
        relatorio += "TEMAS PRIORITÁRIOS (maior impacto de erro)\n"
        relatorio += "-" * 110 + "\n"
        relatorio += (
            f"{'Tema':<45} {'Impacto':>12} {'Respostas Erradas':>18} {'Respostas Avaliadas':>20} {'Taxa de Erro (%)':>16}\n"
        )
        relatorio += "-" * 110 + "\n"

        # ---- Tabela A: Temas (impacto) ----
        if temas:
            temas_imp = []
            for k, e, t in temas:
                p = (float(e) / float(t)) if t else 0.0
                impacto = (float(e) * float(e) / float(t)) if t else 0.0  # penaliza alto erro + grande volume
                temas_imp.append((k, impacto, int(e), int(t), p))
            temas_imp.sort(key=lambda x: x[1], reverse=True)

            for k, imp, e, t, p in temas_imp[:7]:
                relatorio += (
                    f"{str(k)[:45]:<45} {imp:>12.2f} {e:>18d} {t:>20d} {p*100:>15.1f}%\n"
                )
        else:
            relatorio += f"{'Indisponível (dataset sem coluna Tema)':<45} {'-':>12} {'-':>18} {'-':>20} {'-':>16}\n"

        relatorio += "\n"

        # ---- Tabela B: Tendência temporal (médias e variações) ----
        relatorio += "TENDÊNCIA TEMPORAL (médias e variações)\n"
        relatorio += "-" * 80 + "\n"
        relatorio += f"{'Etapa':<15} {'Média':>10} {'Δ Abs':>10} {'Δ %':>10}\n"
        relatorio += "-" * 80 + "\n"

        def _mean_col(col):
            if col and col in df.columns:
                v = pd.to_numeric(df[col], errors="coerce")
                m = float(v.mean())
                return None if math.isnan(m) else m
            return None

        def _fmt(x, nd=1):
            if x is None:
                return "-"
            try:
                return f"{float(x):.{nd}f}"
            except Exception:
                return "-"

        def _fmt_pct(x, nd=1):
            if x is None:
                return "-"
            try:
                return f"{float(x):.{nd}f}%"
            except Exception:
                return "-"

        def _delta(curr, prev):
            if curr is None or prev is None:
                return None, None
            da = curr - prev
            dp = None if abs(prev) < 1e-9 else (da / prev) * 100.0
            return da, dp

        m_pre = _mean_col("Pre_teste")
        m_s1  = _mean_col("Semana1")
        m_s2  = _mean_col("Semana2")
        m_pos = _mean_col("Pos_teste")
        m_ret = _mean_col("Retencao")

        stages = [
            ("Pré", m_pre, None),
            ("1 Sem", m_s1, m_pre),
            ("2 Sem", m_s2, m_s1),
            ("Pós", m_pos, m_s2),
            ("Retenção", m_ret, m_pos),
        ]

        for nome, m_atual, m_prev in stages:
            da, dp = _delta(m_atual, m_prev)
            relatorio += f"{nome:<15} {_fmt(m_atual,1):>10} {_fmt(da,1):>10} {_fmt_pct(dp,1):>10}\n"

        relatorio += "\n"
# interpretação rápida (usa deltas principais)
        d1, _ = _delta(m_s1, m_pre)
        d3, _ = _delta(m_pos, m_s2)
        d4, _ = _delta(m_ret, m_pos)

        interp = []
        if d1 is not None and d1 < 0:
            interp.append("queda inicial (mediação inicial fraca)")
        if d3 is not None and d3 < 0:
            interp.append("queda na consolidação (revisão/avaliação formativa)")
        if d4 is not None and d4 < 0:
            interp.append("perda na retenção (necessário reforço espaçado)")
        if not interp:
            interp.append("tendência sem alertas fortes detectados")
        relatorio += "Leitura rápida: " + "; ".join(interp) + ".\n\n"

# 6) Ranking de risco por aluno (déficit acumulado)
        relatorio += "ALUNOS PRIORITÁRIOS (maior risco por déficit acumulado)\n"
        relatorio += "-" * 100 + "\n"
        relatorio += f"{'Aluno':<25} {'Grupo':<15} {'Risco':>8} {'Déficits (≤40)':<45}\n"
        relatorio += "-" * 100 + "\n"

        cols_notas = [c for c in ["Pre_teste", "Semana1", "Semana2", "Pos_teste", "Retencao"] if c in df.columns]
        col_aluno = "Aluno" if "Aluno" in df.columns else ("Student" if "Student" in df.columns else None)
        col_grupo = "Grupo" if "Grupo" in df.columns else ("Group" if "Group" in df.columns else None)

        riscos = []
        if cols_notas and col_aluno:
            for _, r in df.iterrows():
                aluno = str(r.get(col_aluno, "")).strip()
                if not aluno:
                    continue
                grupo = str(r.get(col_grupo, "")).strip() if col_grupo else ""
                deficit_parts = []
                risco = 0.0
                for c in cols_notas:
                    v = pd.to_numeric(pd.Series([r.get(c)]), errors="coerce").iloc[0]
                    if pd.isna(v):
                        continue
                    d = max(0.0, limiar_baixo - float(v))
                    if d > 0:
                        risco += d
                        # rótulos curtos
                        lab = {"Pre_teste":"Pre", "Semana1":"S1", "Semana2":"S2", "Pos_teste":"Pós", "Retencao":"Ret"}.get(c, c)
                        deficit_parts.append(f"{lab}={float(v):.1f} (déf {d:.1f})")
                if risco > 0:
                    riscos.append((aluno, grupo, risco, "; ".join(deficit_parts)))

            riscos.sort(key=lambda x: x[2], reverse=True)
            for aluno, grupo, risco, det in riscos[:10]:
                relatorio += f"{aluno[:25]:<25} {grupo[:15]:<15} {risco:>8.1f} {det[:45]:<45}\n"
        else:
            relatorio += "Indisponível: coluna de aluno/notas não identificada no dataset.\n"

        relatorio += "\n"

        # ---------------------------------------------------------------------
        # Tabelas detalhadas de erro (Tema / Nível)
        # ---------------------------------------------------------------------
        relatorio += "TOP 7 TEMAS COM MAIOR TAXA DE ERRO\n"
        relatorio += "-" * 100 + "\n"
        if temas:
            relatorio += f"{'Tema':<40} {'Respostas Erradas':>18} {'Respostas Avaliadas':>20} {'Taxa de Erro (%)':>16}\n"
            relatorio += "-" * 100 + "\n"
            for k, e, t in temas[:7]:
                p = (100.0 * float(e) / float(t)) if t else 0.0
                relatorio += f"{str(k)[:40]:<40} {int(e):>18d} {int(t):>20d} {p:>15.1f}%\n"
            relatorio += "\n"
        else:
            relatorio += "Indisponível: não há colunas de questões (Q1..Qn) + Tema no dataset.\n\n"

        relatorio += "ERROS POR NÍVEL COGNITIVO\n"
        relatorio += "-" * 100 + "\n"
        if niveis:
            relatorio += f"{'Nível Cognitivo':<25} {'Respostas Erradas':>18} {'Respostas Avaliadas':>20} {'Taxa de Erro (%)':>16}\n"
            relatorio += "-" * 100 + "\n"
            for k, e, t in niveis:
                p = (100.0 * float(e) / float(t)) if t else 0.0
                relatorio += f"{str(k)[:25]:<25} {int(e):>18d} {int(t):>20d} {p:>15.1f}%\n"
            relatorio += "\n"
        else:
            relatorio += "Indisponível: não há colunas de questões (Q1..Qn) + Nível no dataset.\n\n"

        # ---------------------------------------------------------------------
        # Métricas do modelo (se houver)
        # ---------------------------------------------------------------------
        relatorio += "MÉTRICAS DO MODELO (GLOBAL)\n"
        relatorio += "-" * 100 + "\n"
        if hasattr(self, "metricas") and isinstance(self.metricas, dict) and self.metricas:
            tr = self.metricas.get("train", {})
            va = self.metricas.get("val", {})
            if tr and ("mse" in tr) and ("r2" in tr):
                relatorio += f"Treino    -> MSE: {tr['mse']:.2f} | R²: {tr['r2']:.2f}\n"
            if va and ("mse" in va) and ("r2" in va):
                relatorio += f"Validação -> MSE: {va['mse']:.2f} | R²: {va['r2']:.2f}\n"
        else:
            relatorio += "Indisponível: sem métricas (treine o modelo).\n"
        relatorio += "\n"

        # Métricas por tarefa
        relatorio += "MÉTRICAS POR TAREFA (TREINO)\n"
        relatorio += "-" * 100 + "\n"
        if hasattr(self, "metricas_por_tarefa") and isinstance(self.metricas_por_tarefa, dict) and self.metricas_por_tarefa:
            for tarefa, info in self.metricas_por_tarefa.get("train", {}).items():
                if info and ("n" in info) and ("mse" in info) and ("r2" in info):
                    relatorio += f"• {tarefa:<9} n={info['n']:<4d} MSE={info['mse']:.2f} | R²={info['r2']:.2f}\n"
            relatorio += "\n"
        else:
            relatorio += "• (indisponível)\n\n"

        relatorio += "MÉTRICAS POR TAREFA (VALIDAÇÃO)\n"
        relatorio += "-" * 100 + "\n"
        if hasattr(self, "metricas_por_tarefa") and isinstance(self.metricas_por_tarefa, dict) and self.metricas_por_tarefa:
            for tarefa, info in self.metricas_por_tarefa.get("val", {}).items():
                if info and ("n" in info) and ("mse" in info) and ("r2" in info):
                    relatorio += f"• {tarefa:<9} n={info['n']:<4d} MSE={info['mse']:.2f} | R²={info['r2']:.2f}\n"
            relatorio += "\n"
        else:
            relatorio += "• (indisponível)\n\n"

        relatorio += "=" * 88 + "\n"
        relatorio += "🔍 PREVISÕES (MODELO ÚNICO)\n"
        relatorio += "=" * 88 + "\n"

        # Limpa Tree do relatório
        if False and hasattr(self, "tree_relatorio") and self.tree_relatorio is not None:
            for it in self.tree_relatorio.get_children(""):
                self.tree_relatorio.delete(it)

        if self.modelo is None:
            relatorio += "Indisponível: carregue/treine um modelo.\n\n"
            if hasattr(self, "lbl_relatorio_prev"):
                self.lbl_relatorio_prev.config(text="PREVISÕES — indisponível (treine/carregue um modelo)")
        else:
            dfp = self.dados.copy()

            # Garantir colunas com NaN para facilitar
            for c in ["Pre_teste", "Semana1", "Semana2", "Pos_teste", "Retencao"]:
                if c not in dfp.columns:
                    dfp[c] = np.nan

            # ==========================================================
            # ✅ PREVISÃO SEQUENCIAL (por linha):
            #   Pre+S1           -> prevê S2
            #   Pre+S1+S2        -> prevê Pós
            #   Pre+S1+S2+Pós    -> prevê Retenção
            # ==========================================================
            #
            # Observação importante:
            # Diferente da versão antiga (que escolhia 1 única tarefa para o dataset inteiro),
            # aqui a decisão é feita POR ALUNO, respeitando a disponibilidade real das notas.
            #
            h_pre = "Pre"
            h_s1  = "S1"
            h_s2  = "S2"
            h_pos = "Pós"
            h_ret = "Ret"
            h_etp = "Etapa"
            h_prv = "Prox(Prev)"

            relatorio += f"{'Aluno':<18} {'Grupo':<12} {h_pre:>6} {h_s1:>6} {h_s2:>6} {h_pos:>6} {h_ret:>6} {h_etp:>9} {h_prv:>10}\n"
            relatorio += "-" * 112 + "\n"

            def fmt(v):
                v = pd.to_numeric(pd.Series([v]), errors="coerce").iloc[0]
                if pd.isna(v):
                    return "SN"
                return f"{float(v):.1f}"

            for i in range(len(dfp)):
                ident = self._get_ident_aluno(i)
                row = dfp.iloc[i]
                grupo = str(row.get("Grupo", ""))

                pre = row.get("Pre_teste", np.nan)
                s1  = row.get("Semana1", np.nan)
                s2  = row.get("Semana2", np.nan)
                pos = row.get("Pos_teste", np.nan)
                ret = row.get("Retencao", np.nan)

                # Decide a PRÓXIMA etapa por aluno (sequencial)
                etapa = "---"
                prox_prev = np.nan

                pre_ok = not pd.isna(pd.to_numeric(pre, errors="coerce"))
                s1_ok  = not pd.isna(pd.to_numeric(s1, errors="coerce"))
                s2_ok  = not pd.isna(pd.to_numeric(s2, errors="coerce"))
                pos_ok = not pd.isna(pd.to_numeric(pos, errors="coerce"))
                ret_ok = not pd.isna(pd.to_numeric(ret, errors="coerce"))

                if pre_ok and s1_ok and (not s2_ok):
                    etapa = "S2"
                    prox_prev = self._predict_one_task(pre, s1, np.nan, np.nan, grupo, "Semana2")
                    # opcional: exibir S2 como previsto na própria coluna
                    s2 = prox_prev
                elif pre_ok and s1_ok and s2_ok and (not pos_ok):
                    etapa = "Pós"
                    prox_prev = self._predict_one_task(pre, s1, s2, np.nan, grupo, "Pos_teste")
                    pos = prox_prev
                elif pre_ok and s1_ok and s2_ok and pos_ok and (not ret_ok):
                    etapa = "Ret"
                    prox_prev = self._predict_one_task(pre, s1, s2, pos, grupo, "Retencao")
                    ret = prox_prev

                relatorio += f"{ident:<18} {grupo:<12} {fmt(pre):>6} {fmt(s1):>6} {fmt(s2):>6} {fmt(pos):>6} {fmt(ret):>6} {etapa:>9} {fmt(prox_prev):>10}\n"

            relatorio += "\nLegenda: SN = sem nota (ausente no dataset). Etapa/Prox(Prev) indicam a PRÓXIMA nota prevista por aluno.\n\n"

        relatorio += self._relatorio_baixo_rendimento()

        self.txt_relatorio.config(state=tk.NORMAL)
        self.txt_relatorio.delete(1.0, tk.END)
        self.txt_relatorio.insert(tk.END, relatorio)
        self.txt_relatorio.config(state=tk.DISABLED)

        # Atualiza preview final
        if hasattr(self, "txt_relatorio") and self.txt_relatorio is not None:
            try:
                self.txt_relatorio.config(state=tk.NORMAL)
                self.txt_relatorio.delete(1.0, tk.END)
                self.txt_relatorio.insert(tk.END, relatorio)
                self.txt_relatorio.config(state=tk.DISABLED)
            except Exception:
                pass

        # Salva txt local (mesmo comportamento do relatório)
        try:
            with open("relatorio_desempenho.txt", "w", encoding="utf-8") as f:
                f.write(relatorio)
        except Exception:
            pass

        return relatorio
    def _relatorio_baixo_rendimento(self) -> str:
        ensure_data_libs()
        self._detectar_colunas_curva()
        if self.dados is None or not self.colunas_disponiveis_curva:
            return (
                "\n" + "=" * 88 + "\n"
                f"⚠️ RELATÓRIO — ALUNOS COM BAIXO RENDIMENTO (nota < {int(LIMIAR_BAIXO_RENDIMENTO)})\n"
                + "=" * 88 + "\n"
                "Nenhuma coluna de avaliação disponível para análise.\n"
            )

        cols = [(c, lab) for c, lab in self.colunas_disponiveis_curva if c != "Retencao"]
        if not cols:
            cols = self.colunas_disponiveis_curva[:]

        low_list = []
        for i in range(len(self.dados)):
            row = self.dados.iloc[i]
            deficits = []
            for col, lab in cols:
                v = row.get(col, np.nan)
                try:
                    v = float(v) if pd.notna(v) else np.nan
                except Exception:
                    v = np.nan
                if not np.isnan(v) and v < LIMIAR_BAIXO_RENDIMENTO:
                    deficits.append((LIMIAR_BAIXO_RENDIMENTO - v, lab, v))
            if deficits:
                deficits.sort(reverse=True, key=lambda x: x[0])
                low_list.append((i, str(row.get("Grupo", "")), deficits))

        txt = "\n" + "=" * 88 + "\n"
        txt += f"⚠️ RELATÓRIO — ALUNOS COM BAIXO RENDIMENTO (nota < {int(LIMIAR_BAIXO_RENDIMENTO)})\n"
        txt += "=" * 88 + "\n"

        if not low_list:
            txt += "✅ Nenhum aluno abaixo do limiar nas colunas disponíveis.\n"
            return txt

        count = {}
        pior_por_grupo = {}

        for idx, grupo, deficits in low_list:
            count[grupo] = count.get(grupo, 0) + 1
            pior_nota_aluno = min(d[2] for d in deficits)
            pior_por_grupo[grupo] = min(pior_por_grupo.get(grupo, 1e9), pior_nota_aluno)

        txt += "Resumo por Grupo:\n"
        for grupo in sorted(count.keys()):
            classe = self._classe_por_nota(float(pior_por_grupo.get(grupo, 1e9)))
            txt += f"• {grupo}: {count[grupo]} alunos — {classe}\n"

        txt += "\nLista detalhada:\n"
        for idx, grupo, deficits in low_list:
            ident = self._get_ident_aluno(idx)
            contrib = "; ".join([f"{lab}={v:.1f} (déficit {d:.1f})" for d, lab, v in deficits])
            txt += f"{ident} | Grupo: {grupo} | Contribuições: {contrib}\n"

        txt += "\nAções pedagógicas sugeridas:\n"
        txt += "• Se Pré < limiar: reforçar pré-requisitos e diagnóstico.\n"
        txt += "• Se Semana 1/2 < limiar: aumentar scaffolding, prática guiada e feedback.\n"
        txt += "• Se Pós < limiar: reforçar consolidação, revisão e avaliação formativa.\n"

        return txt


if __name__ == "__main__":
    root = tk.Tk()
    app = EduPredict(root)
    root.mainloop()